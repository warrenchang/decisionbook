from __future__ import annotations

import json, re, shutil, unicodedata, zipfile, html, textwrap, math, os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Iterable, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from markdown_it import MarkdownIt
from jinja2 import Environment, BaseLoader, select_autoescape
from bs4 import BeautifulSoup
import svgwrite
import cairosvg

ROOT = Path('/mnt/data')
SRC_DOCX = ROOT / 'Decision_Persuasion_Negotiation_Textbook.docx'
OUT = ROOT / 'dpn_textbook_github'
ZIP = ROOT / 'Decision_Persuasion_Negotiation_GitHub_Ebook.zip'
EPUB = ROOT / 'Decision_Persuasion_Negotiation_Student_Ebook.epub'

SOURCE_DOCS = [
    ROOT/'01_Why_Decision_Making_Needs_a_Science_Reorganized_Scientifically_Revised_Final(2).docx',
    ROOT/'02_Rational_Choice_Opportunity_Cost_Expected_Utility_Reorganized_Scientifically_Revised_Final(2).docx',
    ROOT/'03 Two Systems of Thinking(1).docx',
    ROOT/'04 Attention, Perception, and Predictive Processing(2).docx',
    ROOT/'05 Valuation, Motivation, and Rationalization(1).docx',
    ROOT/'06 Expectations and Self-Fulfilling Beliefs(2).docx',
    ROOT/'07_Habits_Self_Control_Behavior_Change_Revised_Integrated_Expanded_NoOmissions_Final(2).docx',
    ROOT/'09 Heuristics(1).docx', ROOT/'10 Biases(1).docx', ROOT/'11 Framing Effect(1).docx',
    ROOT/'12 Priming(1).docx', ROOT/'13 Cogntive Ease(1).docx', ROOT/'14 Probability Judgment.docx',
    ROOT/'22 Social Influence(1).docx', ROOT/'25 Persuasion(1).docx',
    ROOT/'26_Storytelling_Narrative_Persuasion_Unabridged_Integrated_v2plus_Final(2).docx',
    ROOT/'27 Communication and Connection(3).docx', ROOT/'28 The nature of negotiation(1).docx',
    ROOT/'29 Distributive Negotiation.docx', ROOT/'30 Integrative Negotiation.docx'
]

# ---------- helpers ----------

def ntext(s: str) -> str:
    s = unicodedata.normalize('NFKC', s or '')
    return re.sub(r'\s+', ' ', s.replace('\u00a0',' ')).strip()

def slugify(s: str) -> str:
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower().replace('&',' and ')
    s = re.sub(r'[^a-z0-9]+','-',s).strip('-')
    return s or 'section'

def esc_md(s: str) -> str:
    s = s.replace('\\','\\\\').replace('|','\\|')
    return s

def normalize_key(s: str) -> str:
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(ch for ch in s if not unicodedata.combining(ch))
    return re.sub(r'[^a-z0-9]+','',s.lower())

@dataclass
class PBlock:
    text: str
    style: str = 'Normal'
    bold: bool = False
    italic: bool = False

@dataclass
class TBlock:
    rows: List[List[str]]
    style: str = ''

Block = Union[PBlock, TBlock]

@dataclass
class Chapter:
    old_num: Optional[int]
    num: int
    title: str
    subtitle: str
    part: int
    blocks: List[Block]
    slug: str = ''
    citations: List[str] = field(default_factory=list)
    references: List[str] = field(default_factory=list)
    figures: List[dict] = field(default_factory=list)


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def paragraph_block(p: Paragraph) -> PBlock:
    text = p.text.replace('\r','').strip()
    # Keep paragraph-level styling if all meaningful runs share it.
    meaningful = [r for r in p.runs if r.text.strip()]
    bold = bool(meaningful) and all(bool(r.bold) for r in meaningful)
    italic = bool(meaningful) and all(bool(r.italic) for r in meaningful)
    return PBlock(text=text, style=p.style.name if p.style else 'Normal', bold=bold, italic=italic)


def table_block(t: Table) -> TBlock:
    rows=[]
    for row in t.rows:
        rows.append([ntext(c.text.replace('\n',' <br> ')) for c in row.cells])
    # remove duplicate cells caused by merged table XML only when exact adjacent repetition is obvious
    return TBlock(rows=rows, style=t.style.name if t.style else '')

# ---------- extract original ----------

def extract_book():
    doc = Document(SRC_DOCX)
    front: List[Block]=[]
    chapters: Dict[int, dict]={}
    appendices: Dict[str,List[Block]]={}
    refs=[]
    current=None; current_appendix=None; pending_num=None
    for obj in iter_blocks(doc):
        if isinstance(obj, Paragraph):
            pb=paragraph_block(obj); txt=ntext(pb.text); st=pb.style
            if st=='Reference':
                if txt: refs.append(txt)
                continue
            if st=='Chapter Label':
                m=re.search(r'(\d+)',txt); pending_num=int(m.group(1)) if m else None
                current={'num':pending_num,'title':'','subtitle':'','blocks':[]}
                chapters[pending_num]=current; current_appendix=None
                continue
            if st=='Chapter Title' and current:
                current['title']=txt; continue
            if st=='Chapter Subtitle' and current:
                current['subtitle']=txt; continue
            if st=='Heading 1' and txt.startswith('Appendix'):
                current=None; current_appendix=txt
                appendices[current_appendix]=[]
                continue
            # ignore old static TOC and course map; they will be regenerated
            if st=='Heading 1' and txt in {'Detailed Table of Contents','Course map','References'}:
                current=None if txt=='References' else current
                current_appendix=None if txt=='References' else current_appendix
                continue
            if st=='TOC Entry':
                continue
            # Ignore old part title: part pages are regenerated.
            if st=='Part Title':
                current=None; current_appendix=None
                continue
            if current is not None:
                current['blocks'].append(pb)
            elif current_appendix:
                appendices[current_appendix].append(pb)
            else:
                front.append(pb)
        else:
            tb=table_block(obj)
            if current is not None: current['blocks'].append(tb)
            elif current_appendix: appendices[current_appendix].append(tb)
            else: front.append(tb)
    return front, chapters, appendices, refs

# ---------- references ----------

def extract_source_references(path: Path) -> List[str]:
    d=Document(path)
    paras=d.paragraphs
    # Find last explicit references heading. If absent, infer final contiguous APA-like run.
    idx=None
    for i,p in enumerate(paras):
        t=ntext(p.text).lower().rstrip(':')
        if t in {'references','reference list','selected references'}:
            idx=i
    candidates=[]
    if idx is not None:
        for p in paras[idx+1:]:
            t=ntext(p.text)
            if t: candidates.append(t)
    else:
        # scan backwards for a run of APA-looking paragraphs
        run=[]
        for p in reversed(paras):
            t=ntext(p.text)
            if not t: continue
            if re.search(r'\((?:18|19|20)\d{2}[a-z]?\)',t) and ('.' in t):
                run.append(t)
            elif len(run)>=5:
                break
        candidates=list(reversed(run))
    # Filter APA references (and remove prompts/headers).
    out=[]
    for t in candidates:
        if re.search(r'\((?:18|19|20)\d{2}[a-z]?\)',t) and len(t)>25:
            out.append(t)
    return out

@dataclass
class Ref:
    text: str
    year: str
    author_text: str
    surnames: List[str]
    first: str
    source: str=''


def parse_ref(text: str, source='') -> Optional[Ref]:
    text=ntext(text)
    m=re.search(r'\((?P<year>(?:18|19|20)\d{2}[a-z]?)\)',text)
    if not m: return None
    year=m.group('year'); author_text=text[:m.start()].strip(' .')
    # author surname extraction: names before initials. Corporate/classical fallbacks.
    surnames=[]
    # Split author string on &, and, and comma boundaries preceding surname+initial pattern.
    # Capture sequences before comma followed by initials.
    pat=re.compile(r'(?:^|,\s*|&\s*|\band\s+)([^,;&]+?),\s*(?:[A-Z][A-Za-z.\- ]*(?:,\s*(?:Jr\.|III))?)')
    for mm in pat.finditer(author_text):
        sn=ntext(mm.group(1)).strip()
        sn=re.sub(r'^(?:&|and)\s+','',sn).strip()
        if sn and len(sn)<80: surnames.append(sn)
    if not surnames:
        if ',' in author_text:
            surnames=[author_text.split(',',1)[0].strip()]
        else:
            surnames=[author_text.split('&',1)[0].strip()]
    return Ref(text=text,year=year,author_text=author_text,surnames=surnames,first=surnames[0],source=source)


def build_reference_db(old_refs: List[str]) -> List[Ref]:
    raw=[(x,'original textbook') for x in old_refs]
    for p in SOURCE_DOCS:
        for x in extract_source_references(p): raw.append((x,p.name))
    # Manual additions used in new synthesis, sourced from the uploaded chapter bibliographies or standard originals.
    manual=[
        'Dahlstrom, M. F. (2014). Using narratives and storytelling to communicate science with nonexpert audiences. Proceedings of the National Academy of Sciences, 111(Suppl. 4), 13614-13620.',
        'Hinyard, L. J., & Kreuter, M. W. (2007). Using narrative communication as a tool for health behavior change: A conceptual, theoretical, and empirical overview. Health Education & Behavior, 34(5), 777-792.',
        'Fisher, R., Ury, W., & Patton, B. (2011). Getting to yes: Negotiating agreement without giving in (3rd ed.). Penguin.',
        'Lax, D. A., & Sebenius, J. K. (1986). The manager as negotiator: Bargaining for cooperation and competitive gain. Free Press.',
        'Bazerman, M. H., & Neale, M. A. (1992). Negotiating rationally. Free Press.',
        'Raiffa, H. (1982). The art and science of negotiation. Harvard University Press.',
        'Raiffa, H., Richardson, J., & Metcalfe, D. (2002). Negotiation analysis: The science and art of collaborative decision making. Belknap Press.',
        'Ames, D. R., & Mason, M. F. (2015). Tandem anchoring: Informational and politeness effects of range offers in social exchange. Journal of Personality and Social Psychology, 108(2), 254-274.',
        'Mason, M. F., Lee, A. J., Wiley, E. A., & Ames, D. R. (2013). Precise offers are potent anchors: Conciliatory counteroffers and attributions of knowledge in negotiations. Journal of Experimental Social Psychology, 49(4), 759-763.',
        'Galinsky, A. D., & Mussweiler, T. (2001). First offers as anchors: The role of perspective-taking and negotiator focus. Journal of Personality and Social Psychology, 81(4), 657-669.',
        'Galinsky, A. D., Maddux, W. W., Gilin, D., & White, J. B. (2008). Why it pays to get inside the head of your opponent: The differential effects of perspective taking and empathy in negotiations. Psychological Science, 19(4), 378-384.',
        'Klein, G. (2007). Performing a project premortem. Harvard Business Review, 85(9), 18-19.',
        'Soll, J. B., Milkman, K. L., & Payne, J. W. (2015). A user’s guide to debiasing. In G. Keren & G. Wu (Eds.), The Wiley Blackwell handbook of judgment and decision making (pp. 924-951). Wiley-Blackwell.',
        'Kahneman, D., Sibony, O., & Sunstein, C. R. (2021). Noise: A flaw in human judgment. Little, Brown Spark.',
        'Edmans, A., García, D., & Norli, Ø. (2007). Sports sentiment and stock returns. Journal of Finance, 62(4), 1967-1998.',
        'Hirshleifer, D., & Shumway, T. (2003). Good day sunshine: Stock returns and the weather. Journal of Finance, 58(3), 1009-1032.',
        'Schwarz, N., & Clore, G. L. (1983). Mood, misattribution, and judgments of well-being: Informative and directive functions of affective states. Journal of Personality and Social Psychology, 45(3), 513-523.',
        'Taylor, S. E., & Brown, J. D. (1988). Illusion and well-being: A social psychological perspective on mental health. Psychological Bulletin, 103(2), 193-210.',
        'Henrich, J. (2016). The secret of our success: How culture is driving human evolution, domesticating our species, and making us smarter. Princeton University Press.',
        'Horner, V., & Whiten, A. (2005). Causal knowledge and imitation/emulation switching in chimpanzees (Pan troglodytes) and children (Homo sapiens). Animal Cognition, 8(3), 164-181.',
        'van Baaren, R. B., Holland, R. W., Kawakami, K., & van Knippenberg, A. (2004). Mimicry and prosocial behavior. Psychological Science, 15(1), 71-74.',
        'Greenwald, A. G. (1980). The totalitarian ego: Fabrication and revision of personal history. American Psychologist, 35(7), 603-618.',
        'Heider, F. (1958). The psychology of interpersonal relations. Wiley.',
        'Kelley, H. H. (1967). Attribution theory in social psychology. In D. Levine (Ed.), Nebraska symposium on motivation (Vol. 15, pp. 192-238). University of Nebraska Press.',
        'Festinger, L. (1957). A theory of cognitive dissonance. Stanford University Press.',
        'Tajfel, H., Billig, M. G., Bundy, R. P., & Flament, C. (1971). Social categorization and intergroup behaviour. European Journal of Social Psychology, 1(2), 149-178.',
        'Olson, R. (2015). Houston, we have a narrative: Why science needs story. University of Chicago Press.',
        'Heath, C., & Heath, D. (2007). Made to stick: Why some ideas survive and others die. Random House.',
        'Thibodeau, P. H., & Boroditsky, L. (2011). Metaphors we think with: The role of metaphor in reasoning. PLOS ONE, 6(2), e16782.',
        'Hekkert, P. P. M., Snelders, H. M. J. J., & van Wieringen, P. C. W. (2003). Most advanced, yet acceptable: Typicality and novelty as joint predictors of aesthetic preference in industrial design. British Journal of Psychology, 94(1), 111-124.',
    ]
    raw += [(x,'manual consolidation') for x in manual]
    by_norm={}
    for t,src in raw:
        t=t.replace('Wood, W., & Ruenger, D. (2016). Psychology of habit.', 'Wood, W., & Rünger, D. (2016). Psychology of habit.')
        if t.startswith('Simon, H. A. (1955). A behavioral model of rational choice.'):
            t='Simon, H. A. (1955). A behavioral model of rational choice. The Quarterly Journal of Economics, 69(1), 99-118.'
        if t.startswith('Aristotle. (2007). On rhetoric: A theory of civic discourse'):
            t='Aristotle. (2007). On rhetoric: A theory of civic discourse (G. A. Kennedy, Trans., 2nd ed.). Oxford University Press. (Original work ca. 4th century BCE).'
        if t.startswith('Dahlstrom, M. F. (2014). Using narratives and storytelling to communicate science with nonexpert audiences.'):
            t='Dahlstrom, M. F. (2014). Using narratives and storytelling to communicate science with nonexpert audiences. Proceedings of the National Academy of Sciences, 111(Suppl. 4), 13614-13620.'
        if t.startswith('Raiffa, H., Richardson, J., & Metcalfe, D. (2002). Negotiation analysis:'):
            t='Raiffa, H., Richardson, J., & Metcalfe, D. (2002). Negotiation analysis: The science and art of collaborative decision making. Belknap Press.'
        if t.startswith('Nisbett, R. E., & Wilson, T. D. (1977). The halo effect:'):
            t=t.replace('(1977).','(1977a).',1)
        if t.startswith('Nisbett, R. E., & Wilson, T. D. (1977). Telling more than we can know:'):
            t=t.replace('(1977).','(1977b).',1)
        r=parse_ref(t,src)
        if not r: continue
        k=normalize_key(r.text)
        # Prefer original textbook wording; otherwise first source.
        if k not in by_norm or src=='original textbook': by_norm[k]=r
    return list(by_norm.values())


def citation_patterns(r: Ref):
    y=re.escape(r.year)
    ss=[re.escape(s).replace(r'\ ',r'\s+') for s in r.surnames]
    first=ss[0]
    pats=[]
    if len(ss)==1:
        pats += [rf'\b{first}(?:[’\']s)?\s*\(\s*{y}\s*\)', rf'(?<![A-Za-z]){first}\s*,\s*{y}\b']
    elif len(ss)==2:
        a,b=ss[:2]
        pats += [rf'\b{a}\s+(?:&|and)\s+{b}(?:[’\']s)?\s*\(\s*{y}\s*\)',
                 rf'(?<![A-Za-z]){a}\s*(?:&|and)\s*{b}\s*,\s*{y}\b',
                 rf'\b{a}\s+et\s+al\.?(?:[’\']s)?\s*\(\s*{y}\s*\)', rf'(?<![A-Za-z]){a}\s+et\s+al\.?,\s*{y}\b']
    else:
        pats += [rf'\b{first}\s+et\s+al\.?(?:[’\']s)?\s*\(\s*{y}\s*\)',rf'(?<![A-Za-z]){first}\s+et\s+al\.?,\s*{y}\b']
        # Full author list in parenthetical citations, allowing punctuation.
        seq=first
        for j,b in enumerate(ss[1:]):
            if j==len(ss)-2: seq += rf'\s*,?\s*(?:&|and)\s*{b}'
            else: seq += rf'\s*,\s*{b}'
        pats += [rf'\b{seq}(?:[’\']s)?\s*\(\s*{y}\s*\)',rf'(?<![A-Za-z]){seq}\s*,\s*{y}\b']
    # APA permits original/reprint citations such as 1934/2010. The reference
    # entry carries the reprint year and an original-work note.
    om=re.search(r'Original work (?:published|ca\.)\s*((?:18|19|20)\d{2})',r.text,re.I)
    if om:
        oy=re.escape(om.group(1))
        if len(ss)==1:
            pats += [rf'\b{first}(?:[’\']s)?\s*\(\s*{oy}\s*/\s*{y}\s*\)',
                     rf'(?<![A-Za-z]){first}\s*,\s*{oy}\s*/\s*{y}\b']
        elif len(ss)==2:
            a,b=ss[:2]
            pats += [rf'\b{a}\s+(?:&|and)\s+{b}(?:[’\']s)?\s*\(\s*{oy}\s*/\s*{y}\s*\)',
                     rf'(?<![A-Za-z]){a}\s*(?:&|and)\s*{b}\s*,\s*{oy}\s*/\s*{y}\b']
        else:
            pats += [rf'\b{first}\s+et\s+al\.?(?:[’\']s)?\s*\(\s*{oy}\s*/\s*{y}\s*\)',
                     rf'(?<![A-Za-z]){first}\s+et\s+al\.?,\s*{oy}\s*/\s*{y}\b']
    # APA compresses repeated authors into one citation, e.g. Simon (1955,
    # 1956) or (Berridge & Robinson, 1998, 2003). Match the target year
    # anywhere in the same author segment, without crossing a semicolon.
    bases=[]
    if len(ss)==1:
        bases=[first]
    elif len(ss)==2:
        a,b=ss[:2]; bases=[rf'{a}\s+(?:&|and)\s+{b}',rf'{a}\s+et\s+al\.?']
    else:
        bases=[rf'{first}\s+et\s+al\.?',seq]
    for base in bases:
        pats += [rf'\b{base}(?:[’\']s)?\s*\((?=[^)]*\b{y}\b)[^)]*\)',
                 rf'(?<![A-Za-z]){base}\s*,(?=[^;)]*\b{y}\b)[^;)]*']
    return pats


def refs_for_text(text: str, refdb: List[Ref]) -> List[Ref]:
    # Fast candidate filter first; precise regex only on likely surname-year pairs.
    low=unicodedata.normalize("NFKC",text).lower()
    found=[]
    for r in refdb:
        if r.year.lower() not in low or r.first.lower() not in low:
            continue
        if any(re.search(p,text,re.I|re.S) for p in citation_patterns(r)):
            found.append(r)
    seen=set(); out=[]
    for r in sorted(found,key=lambda z:(normalize_key(z.author_text),z.year,normalize_key(z.text))):
        k=normalize_key(r.text)
        if k not in seen: seen.add(k); out.append(r)
    return out

# ---------- chapter restructuring ----------

PARTS={
1:('How Choices Take Shape','From the hidden process to predictions, values, and expectations.'),
2:('Judgment Under Uncertainty and Context','How shortcuts, statistics, frames, and mental context bend belief.'),
3:('From Choice to Habit and Learning','How explanations, repetition, craving, and design shape behavior over time.'),
4:('The Social Mind','How other people become evidence, audiences, norms, and sources of authority.'),
5:('Persuasion, Story, and Connection','How messages move models and how conversation creates shared understanding.'),
6:('Negotiation','How interdependent decision-makers claim, create, and implement value.'),
7:('Better Judgment in Practice','How to build repeatable decision hygiene across the course.'),
}

ORDER=[
 (1,1),(2,1),(3,1),(4,1),(5,1),(6,1),(7,1),(8,1),
 (10,2),(11,2),(12,2),(18,2),(13,2),(14,2),(15,2),(16,2),(17,2),
 (9,3),(19,3),(20,3),(21,3),
 ('22a',4),('23a',4),('23b',4),
 (24,5),(25,5),(26,5),(27,5),(28,5),
 (29,6),(30,6),(31,6),(32,6),(33,6),
 (34,7)
]


def find_heading_idx(blocks: List[Block], heading: str) -> int:
    for i,b in enumerate(blocks):
        if isinstance(b,PBlock) and b.style=='Heading 2' and ntext(b.text).lower()==heading.lower(): return i
    return -1


def strip_endmatter(blocks: List[Block]) -> List[Block]:
    for h in ('Key ideas','Study and practice'):
        idx=find_heading_idx(blocks,h)
        if idx>=0: return blocks[:idx]
    return blocks


def custom_endmatter(key_ideas: List[str], questions: List[str], activities: List[str]) -> List[Block]:
    out=[PBlock('Key ideas','Heading 2')]
    out += [PBlock(x,'Takeaway Bullet') for x in key_ideas]
    out.append(PBlock('Study and practice','Heading 2'))
    out.append(PBlock('Questions','Heading 3'))
    out += [PBlock(f'{i}. {q}','Normal') for i,q in enumerate(questions,1)]
    out.append(PBlock('Activities','Heading 3'))
    out += [PBlock(f'{i}. {q}','Normal') for i,q in enumerate(activities,1)]
    out.append(PBlock('Experience-Explain-Apply reflection','Heading 3'))
    out.append(PBlock('In no more than 120 words: describe one specific experience, explain it using one or two concepts from the chapter, and state one application or testable prediction.','Normal'))
    return out


def make_social_chapters(ch22, ch23):
    b22=ch22['blocks']; b23=ch23['blocks']
    i_conf=find_heading_idx(b22,'Conformity and the pull of the group')
    b22core=strip_endmatter(b22[:i_conf] if i_conf>=0 else b22)
    conf_section=[]
    if i_conf>=0:
        end=find_heading_idx(b22,'Key ideas')
        conf_section=b22[i_conf:end if end>=0 else len(b22)]
    i_auth=find_heading_idx(b23,'Authority and responsibility')
    i_key=find_heading_idx(b23,'Key ideas')
    # Norm/social proof material from start through before authority.
    b23norm=strip_endmatter(b23[:i_auth] if i_auth>=0 else b23)
    b23auth=b23[i_auth:i_key if i_key>=0 else len(b23)] if i_auth>=0 else []

    c22={
      'title':'Social Learning, Mimicry, and Attribution',
      'subtitle':'How other people become models, mirrors, and causal explanations',
      'blocks':b22core + custom_endmatter(
        ['Human intelligence is culturally cumulative because people learn from other people.',
         'Imitation and mimicry can coordinate interaction, but prestige and similarity are imperfect signals of expertise.',
         'Attributions determine whether we punish, coach, ask, trust, or renegotiate.',
         'Perspective-getting is usually safer than confident mind-reading.'],
        ['Why is social learning both a human advantage and a source of systematic error?','How does attribution change the next action in a conflict?','When is mimicry affiliative, and when could it become manipulative?'],
        ['Attribution flip: replace one recent character judgment with two situational hypotheses and one question.','Social-learning audit: identify whose behavior you treated as evidence and whether their judgment was independent.'])
    }
    c23={
      'title':'Conformity, Norms, and Social Proof',
      'subtitle':'How groups define what feels true, normal, safe, and desirable',
      'blocks':[
        PBlock('Learning goals','Heading 2'),
        PBlock('Distinguish informational from normative influence.','Learning Objective'),
        PBlock('Explain how conformity, descriptive norms, injunctive norms, and social proof shape behavior.','Learning Objective'),
        PBlock('Recognize cascades, pluralistic ignorance, boomerang effects, and the role of dissent.','Learning Objective'),
        PBlock('Design truthful norm messages that do not normalize the behavior they seek to reduce.','Learning Objective'),
        TBlock([['CORE IDEA'],['Groups are not merely external pressure. They alter what counts as evidence, what seems normal, and what disagreement appears to cost.']])
      ] + conf_section + b23norm + custom_endmatter(
        ['Conformity can reflect a desire to be correct, accepted, or both.','Descriptive norms say what people do; injunctive norms say what people approve.','Popularity is informative only to the extent that judgments are independent and the signal is not manipulated.','A single ally can transform the social cost of dissent.'],
        ['Why can broadcasting a high rate of undesirable behavior increase it?','When is the crowd good evidence, and when is it a hall of mirrors?','How should a leader make private doubt speakable?'],
        ['Norm-message lab: write descriptive and injunctive versions of one message, then identify the backfire risk.','Dissent design: redesign a meeting so the first honest disagreement does not require heroism.'])
    }
    c24={
      'title':'Authority, Bystanders, and Influence Triggers',
      'subtitle':'Signals that coordinate behavior—and can be counterfeited',
      'blocks':[
        PBlock('Learning goals','Heading 2'),
        PBlock('Explain authority, diffusion of responsibility, pluralistic ignorance, and bystander intervention.','Learning Objective'),
        PBlock('Analyze reciprocity, consistency, scarcity, liking, unity, social proof, and reasons as recurring influence triggers.','Learning Objective'),
        PBlock('Distinguish legitimate signals from counterfeit cues.','Learning Objective'),
        PBlock('Design and resist influence while preserving truth, welfare, and agency.','Learning Objective'),
        TBlock([['CORE IDEA'],['Influence triggers are useful because they compress social information. They become manipulative when the cue is manufactured, material facts are hidden, or exit is made difficult.']])
      ] + b23auth + custom_endmatter(
        ['Authority is a coordination device, not a substitute for moral responsibility.','Bystander inaction often reflects failures of noticing, interpretation, and assigned responsibility.','Reciprocity, consistency, scarcity, liking, unity, and social proof are cue-based shortcuts that can be legitimate or counterfeited.','Ethical influence makes the mechanism inspectable and preserves a meaningful path to refuse.'],
        ['Why can ordinary people comply with harmful authority?','What makes a scarcity signal informative rather than manipulative?','How can an emergency system assign responsibility before ambiguity arrives?'],
        ['Influence audit: analyze a recent request and identify the cue, the information it legitimately conveyed, and the information it concealed.','Bystander redesign: turn an ambiguous request for help into an assigned, observable action.'])
    }
    return c22,c23,c24

# Citation insertions by new chapter number. Inserted after learning goals/core box.
CITE_INSERTS={
1:[('The distinction among normative, descriptive, and prescriptive decision science follows a long tradition of formal decision theory, bounded rationality, and decision-improvement research (Edwards, 1954; Simon, 1955; Milkman et al., 2009).','The hidden process'),
   ('Outcome bias and hindsight bias show why a good process cannot be reconstructed reliably from the result alone (Baron & Hershey, 1988; Fischhoff, 1975).','Outcome is not process')],
2:[('Expected utility and subjective expected utility are normative representations of coherent choice under specified assumptions, not descriptions of a conscious mental algorithm (von Neumann & Morgenstern, 1944; Savage, 1954; Edwards, 1954).','A family of models'),
   ('Bounded rationality and satisficing recognize that search, information, and computation have costs (Simon, 1955, 1956).','Bounded rationality and satisficing')],
3:[('Opportunity costs are often neglected when the best forgone alternative is not explicitly represented (Frederick et al., 2009).','Information has value only when it can change action'),
   ('Value-of-information and multi-attribute analysis make the benefits of learning and trade-offs inspectable (Raiffa, 1968; Keeney & Raiffa, 1976).','Multiple objectives and decision tables')],
9:[('Heuristics are adaptive strategies for bounded agents; their quality depends on the fit between a rule and the environment in which it is used (Simon, 1955, 1956; Gigerenzer & Gaissmaier, 2011).','Trigger features and cue-based judgment')],
10:[('Availability and affect can be sensible cues, but vividness, recency, and incidental emotion can make them diverge from frequency and diagnostic evidence (Tversky & Kahneman, 1973; Slovic et al., 2002; Lerner et al., 2015).','Availability: what comes to mind feels common')],
11:[('Representativeness substitutes resemblance for probability, which helps explain base-rate neglect and conjunction errors (Kahneman & Tversky, 1972; Tversky & Kahneman, 1983).','The representative story')],
12:[('Natural-frequency formats, reference classes, and calibration make statistical structure easier to see and improve probabilistic judgment (Gigerenzer & Hoffrage, 1995; Tetlock & Gardner, 2015).','Conditional probability and the base-rate trap')],
18:[('People can give sincere explanations for choices without having reliable introspective access to the processes that produced them (Nisbett & Wilson, 1977b; Johansson et al., 2005).','Reasons are not the same as causes')],
19:[('Habits are learned context-response associations that strengthen through repetition in stable settings; frequency alone is not the definition of habit (Wood et al., 2002; Wood & Neal, 2007; Wood & Rünger, 2016).','How habits form'),
    ('Automaticity develops gradually and varies substantially by behavior and person rather than following a fixed 21-day rule (Lally et al., 2010).','How habits form')],
20:[('Wanting, liking, and choosing depend on partly separable processes; dopamine is more closely tied to learning, incentive salience, and pursuit than to pleasure itself (Berridge & Robinson, 1998, 2003; Schultz et al., 1997).','Wanting is not liking'),
    ('Beneficial self-control often operates upstream through situation selection and reliable habits rather than constant resistance at the moment of temptation (Galla & Duckworth, 2015; Duckworth et al., 2016).','Self-control without heroics')],
21:[('Implementation intentions strengthen the link between a specified cue and a planned response, especially when paired with a meaningful goal (Gollwitzer, 1999; Gollwitzer & Sheeran, 2006).','Implementation intentions'),
    ('Mental contrasting, friction, prompts, defaults, and context change influence whether motivation becomes behavior (Oettingen, 2012; Johnson & Goldstein, 2003; Dai et al., 2014).','Choosing the right tool')],
25:[('Persuasion depends on both message arguments and the audience’s motivation and ability to elaborate; source cues matter differently under different processing conditions (Petty & Cacioppo, 1986; Chaiken, 1980).','Two routes through a message')],
26:[('Narratives organize experience through agents, intentions, conflict, and consequence, complementing abstract or paradigmatic reasoning (Bruner, 1986, 1990; Fisher, 1984, 1987).','What is a story?'),
    ('Transportation, identification, and mental simulation can reduce counterarguing and make consequences experientially meaningful (Green & Brock, 2000; Cohen, 2001; van Laer et al., 2014).','The psychology of narrative persuasion')],
27:[('Narrative evidence is most defensible when a case is connected to representative data rather than allowed to substitute for base rates (Small et al., 2007; Slovic, 2007; Dahlstrom, 2014).','The story-evidence braid')],
30:[('Negotiation is a joint decision process in which communication, alternatives, beliefs, preferences, and strategic interdependence jointly determine the agreement set (Raiffa, 1982; Lax & Sebenius, 1986; Thompson, 2015).','What negotiation is not')],
31:[('A negotiator’s BATNA and reservation value discipline acceptance decisions, while ambitious but defensible aspirations can improve bargaining outcomes (Fisher et al., 2011; White & Neale, 1994).','BATNA and bargaining power')],
32:[('First offers can anchor final settlements, but an uninformed first offer can reveal ignorance as well as shape the range (Galinsky & Mussweiler, 2001).','Responding to anchors'),
    ('Precise and range offers affect perceived knowledge and counteroffers, but their credibility depends on context and justification (Mason et al., 2013; Ames & Mason, 2015).','Responding to anchors')],
33:[('Fixed-pie beliefs and inaccurate social perception prevent negotiators from discovering compatible interests and differences in priorities (Thompson & Hastie, 1990).','Barriers to value creation'),
    ('Perspective-taking can improve discovery of integrative agreements, particularly when it supports diagnostic questioning rather than untested projection (Galinsky et al., 2008).','Interests behind positions')],
34:[('MESOs and contingent agreements help reveal priorities and manage differences in beliefs, provided outcomes are measurable and incentives are well aligned (Raiffa et al., 2002; Thompson, 2015).','Contingent contracts')],
35:[('Decision hygiene improves judgment by changing the process around the judge: independent estimates, structured aggregation, premortems, calibration, and after-action review (Klein, 2007; Milkman et al., 2009; Soll et al., 2015; Kahneman et al., 2021).','Before the decision: improve the input')]
}

CROSSREF_REPL={
 'The four chapters follow that hidden sequence.':'The chapters follow that hidden sequence.',
 'another chapter mapped the full decision loop':'Chapter 1 mapped the full decision loop',
 'another chapter introduced the rational choice model':'Chapter 2 introduced the rational choice model',
 'another chapter emphasized the distinction between prediction and valuation':'Chapter 1 emphasized the distinction between prediction and valuation',
 'The rational choice model described in another chapter':'The rational-choice benchmark in Chapter 2',
 'the representativeness heuristic, introduced in another chapter':'the representativeness heuristic, introduced in Chapter 11',
 'another chapter described rationalization':'Chapter 18 described rationalization',
 'the later chapter on rationalization':'Chapter 18 on rationalization',
 'a later chapter turns to risk and expected utility':'Chapter 2 explains expected utility under risk and uncertainty',
 'A later chapter turns to culture.':'The broader cultural implications appear throughout the social and negotiation parts.',
 'the chapter on prospect theory':'Chapter 15 on framing and reference points',
 'the chapter on social preferences':'the social-mind and negotiation chapters',
 'The next chapter turns to nudge and choice architecture.':'Chapter 21 showed how choice architecture and friction design structure behavior.',
 'next chapter turns to nudge and choice architecture':'Chapter 21 showed how choice architecture and friction design structure behavior',
 'A later chapter develops this point through feasible sets and opportunity cost.':'Chapter 3 develops this point through feasible sets and opportunity cost.',
 'another chapter asks a narrower normative question':'Chapter 2 asks a narrower normative question',
 'the rational choice model from another chapter':'the rational-choice benchmark from Chapter 2',
 'another chapter introduced the distinction between prediction and valuation':'Chapter 1 introduced the distinction between prediction and valuation',
 'That is the subject of a later chapter.':'Chapter 8 examines how expectations can become part of the causal system.',
 'In another chapter terms':'In Chapter 6’s terms',
 'A later chapter turns to habits.':'Part III turns from explanation to habits and behavior design.',
 'another chapter turns next to biases.':'Chapters 10–17 examine the shortcuts, statistical errors, and contextual mechanisms that can bend judgment.',
 'The representativeness heuristic, introduced in another chapter':'The representativeness heuristic, introduced in Chapter 9',
 'discussed in another chapter':'discussed in Chapter 13',
 'A later chapter turns to risk and expected utility.':'Chapter 2 showed how probabilities and utilities are combined under risk; this chapter supplies the statistical discipline those models require.',
 'which will be discussed in a later chapter':'which is central to prospect theory and reference-dependent choice',
 'Priming, discussed in a later chapter':'Priming, discussed in Chapter 16',
 'This pattern will return in Chapter 15 on framing and reference points.':'This gain–loss asymmetry is one of prospect theory’s central patterns.',
 'another chapter’s discussion of predictive processing':'Chapter 6’s discussion of predictive processing',
 'which will be a later chapter':'which is examined in Chapter 17',
 'A later chapter turns to cognitive ease, fluency, and familiarity.':'Chapter 17 turns to cognitive ease, fluency, and familiarity.',
 'A later chapter turns to rationalization and the narrating self.':'Chapter 18 showed how the narrating self can rationalize actions after they occur.',
 'A later chapter turns to culture, identity, and social meaning.':'The broader cultural layer shapes the shared meanings within which immediate social influence operates.',
 'A later chapter on storytelling will examine this more deeply.':'Chapters 26 and 27 examine how stories create meaning and how they should be braided with evidence.',
 'A later chapter turns to storytelling.':'Chapters 26 and 27 turn to storytelling and evidence-aligned narrative design.',
 'Value claiming will be developed in another chapter. Value creation will be developed in another chapter.':'Chapters 31 and 32 develop value claiming; Chapters 33 and 34 develop value creation and advanced agreement design.',
 'As another chapter discussed, face':'As Chapter 29 discussed, face',
 'another chapter warned that ethical fading':'Chapter 13 warned that ethical fading',
 'the preparation discussed in another chapter':'the preparation discussed in Chapter 31',
 'another chapter distinguished perspective-taking from perspective-getting':'Chapter 28 distinguished perspective-taking from perspective-getting',
}


def patch_text(text: str) -> str:
    t=text
    for a,b in CROSSREF_REPL.items(): t=t.replace(a,b)
    t=t.replace('Confirmation, self-service, and the feeling of knowing','Confirmation, self-serving judgment, and the feeling of knowing')
    # Generic cleanup of a few inaccurate forward references.
    t=t.replace('In a later chapter, the narrating self will be discussed','Chapter 18 discussed the narrating self')
    t=t.replace('The later chapter on social influence','The social-influence chapters')
    return t


def clean_chapter_flow(ch: Chapter):
    # A few source transitions referred to the former chapter order. Rewrite
    # them as explicit bridges in the new cumulative structure.
    flow_replacements={
      4:{
        'This rule also connects another chapter to the rest of the book. another chapter will show that attention and perception shape what System 1 offers to consciousness. another chapter will examine valuation, motivation, habit, and rationalization. Chapters 6 and 7 will explore how fast thinking produces heuristics and biases. Chapters 8 and 9 will show how frames and fluency shape what feels right. another chapter will show why probability often requires slow correction. Later chapters will apply these ideas to risk, time, finance, social influence, persuasion, strategy, and negotiation.':
        'This framework connects Chapter 4 to the rest of the book. Chapter 5 examines what attention admits as evidence; Chapters 6–8 explain prediction, valuation, and expectation; Chapters 9–17 examine shortcuts, probability, bias, framing, priming, and fluency; Parts IV–VI apply these mechanisms to influence, communication, and negotiation.'
      },
      7:{
        'The chapter ends with a bridge. Valuation tells us what feels worth doing. But when a valuation is repeatedly attached to the same cue and action, the decision may stop feeling like a decision. It becomes automatic. That is why habit deserves its own chapter.':
        'The chapter ends with two bridges. Chapter 8 first examines expectations: beliefs that can change attention, action, and feedback. Part III then follows repeated cue–value–action loops as they become habits.'
      },
      14:{
        'Nisbett and Wilson (1977) demonstrated a halo effect':'Nisbett and Wilson (1977a) demonstrated a halo effect'
      },
      18:{
        '(Nisbett & Wilson, 1977)':'(Nisbett & Wilson, 1977b)',
        'Nisbett and Wilson (1977) reviewed':'Nisbett and Wilson (1977b) reviewed'
      },
      21:{
        'The DPN habit map in Figure 7.3 translates the same habit logic across three domains. Decision habits protect thought. Persuasion habits design transparent paths. Negotiation habits prepare behavior for stress.':
        'The same habit logic operates across the course: decision habits protect thought, persuasion habits design transparent paths, and negotiation habits prepare behavior for stress.',
      },
      34:{
        'A later chapter turns to nudge and choice architecture. The connection is natural. Negotiation designs agreements between parties. Choice architecture designs environments in which people make decisions. Both ask how structure changes behavior. Both require attention to information, incentives, defaults, friction, timing, salience, social meaning, and ethics. A good negotiator designs a better conversation. A good choice architect designs a better decision environment.':
        'Chapter 35 turns from individual agreements to recurring decision systems. Negotiation designs a better conversation and agreement; decision hygiene designs the agendas, roles, forecasts, checks, and feedback through which future decisions will be made.'
      }
    }
    repls=flow_replacements.get(ch.num,{})
    for b in ch.blocks:
        if isinstance(b,PBlock):
            for a,z in repls.items(): b.text=b.text.replace(a,z)
        else:
            b.rows=[[cell.replace(a,z) for a,z in repls.items()] for row in b.rows for cell in []] if False else b.rows

    # Convert the predictive-loop matrix from an orphan figure caption to a table title.
    if ch.num==6:
        for b in ch.blocks:
            if isinstance(b,PBlock):
                b.text=b.text.replace('Figure 6.1  A functional predictive loop','Table 6.1. A functional predictive loop')

    # Remove captions for source figures that were not embedded in the DOCX;
    # the online edition supplies new accessible diagrams instead.
    kept=[]
    for b in ch.blocks:
        if isinstance(b,PBlock) and re.match(r'^Figure\s+\d+\.\d+\.?(?:\s|$)',ntext(b.text),re.I):
            continue
        kept.append(b)
    ch.blocks=kept

    # Renumber retained table titles to match the revised chapter sequence.
    table_ids=[]
    for b in ch.blocks:
        texts=[b.text] if isinstance(b,PBlock) else sum(b.rows,[])
        for t in texts:
            for m in re.finditer(r'\bTable\s+(\d+\.\d+)\b',t,re.I):
                if m.group(1) not in table_ids: table_ids.append(m.group(1))
    tmap={old:f'{ch.num}.{i+1}' for i,old in enumerate(table_ids)}
    for b in ch.blocks:
        if isinstance(b,PBlock):
            for old,new in tmap.items(): b.text=re.sub(rf'\bTable\s+{re.escape(old)}\b',f'Table {new}',b.text,flags=re.I)
        else:
            for ri,row in enumerate(b.rows):
                for ci,cell in enumerate(row):
                    for old,new in tmap.items(): cell=re.sub(rf'\bTable\s+{re.escape(old)}\b',f'Table {new}',cell,flags=re.I)
                    b.rows[ri][ci]=cell


def insert_citation_paragraphs(ch: Chapter):
    inserts=CITE_INSERTS.get(ch.num,[])
    if not inserts: return
    # Insert immediately after matching heading's following first content block, or before key ideas if no match.
    for text,heading in inserts:
        idx=-1
        for i,b in enumerate(ch.blocks):
            if isinstance(b,PBlock) and b.style=='Heading 2' and heading.lower() in ntext(b.text).lower(): idx=i
        if idx<0:
            idx=next((i for i,b in enumerate(ch.blocks) if isinstance(b,PBlock) and b.style=='Heading 2' and b.text=='Key ideas'),len(ch.blocks))
        # Insert after heading and any immediate core callout, but before next heading.
        pos=min(idx+2,len(ch.blocks))
        ch.blocks.insert(pos,PBlock(text,'Normal'))


def restructure(front, raw_chapters):
    c22,c23,c24=make_social_chapters(raw_chapters[22],raw_chapters[23])
    special={'22a':c22,'23a':c23,'23b':c24}
    out=[]
    newnum=1
    for old,part in ORDER:
        data=special[old] if isinstance(old,str) else raw_chapters[old]
        title=data['title']; subtitle=data['subtitle']; blocks=data['blocks'][:]
        ch=Chapter(old_num=None if isinstance(old,str) else old,num=newnum,title=title,subtitle=subtitle,part=part,blocks=blocks)
        ch.slug=f'{newnum:02d}-{slugify(title)}'
        # Patch text and chapter numbers in all blocks.
        for b in ch.blocks:
            if isinstance(b,PBlock): b.text=patch_text(b.text)
            else: b.rows=[[patch_text(cell) for cell in row] for row in b.rows]
        clean_chapter_flow(ch)
        insert_citation_paragraphs(ch)
        out.append(ch); newnum+=1
    return out

# ---------- figures ----------

COL={'navy':'#17324d','blue':'#2f6f9f','light':'#eaf3f8','teal':'#2b7a78','orange':'#c56a2d','sand':'#f7f1e8','gray':'#66737f','white':'#ffffff','red':'#a74343','green':'#447c55'}

def new_svg(path: Path, w=1200,h=650,title='',desc=''):
    dwg=svgwrite.Drawing(str(path),size=(w,h),viewBox=f'0 0 {w} {h}')
    dwg.add(svgwrite.base.Title(title)); dwg.add(svgwrite.base.Desc(desc))
    dwg.add(dwg.rect(insert=(0,0),size=(w,h),fill='white'))
    return dwg

def txt(dwg,x,y,s,size=28,weight='normal',anchor='middle',fill=None):
    dwg.add(dwg.text(s,insert=(x,y),font_family='Arial, sans-serif',font_size=size,font_weight=weight,text_anchor=anchor,fill=fill or COL['navy']))

def box(dwg,x,y,w,h,label,sub='',fill=None,stroke=None,round=16):
    dwg.add(dwg.rect(insert=(x,y),size=(w,h),rx=round,ry=round,fill=fill or COL['light'],stroke=stroke or COL['blue'],stroke_width=3))
    txt(dwg,x+w/2,y+h/2-8,label,26,'bold')
    if sub: txt(dwg,x+w/2,y+h/2+26,sub,17,'normal',fill=COL['gray'])

def arrow(dwg,x1,y1,x2,y2,stroke=None,width=4):
    stroke=stroke or COL['navy']
    dwg.add(dwg.line(start=(x1,y1),end=(x2,y2),stroke=stroke,stroke_width=width))
    ang=math.atan2(y2-y1,x2-x1); L=15
    p1=(x2-L*math.cos(ang-math.pi/6),y2-L*math.sin(ang-math.pi/6))
    p2=(x2-L*math.cos(ang+math.pi/6),y2-L*math.sin(ang+math.pi/6))
    dwg.add(dwg.polygon(points=[(x2,y2),p1,p2],fill=stroke))

def save_svg(dwg,path):
    dwg.save(pretty=True)
    cairosvg.svg2png(url=str(path),write_to=str(path.with_suffix('.png')),output_width=1800)


def generate_figures(figdir: Path):
    figdir.mkdir(parents=True,exist_ok=True)
    figs={}
    # 1 decision loop
    p=figdir/'decision-loop.svg'; d=new_svg(p,title='The decision loop',desc='Alternatives and information feed prediction and valuation, which guide choice; outcome and feedback update the loop.')
    txt(d,600,48,'The decision loop',34,'bold')
    coords=[(65,255,200,100,'Alternatives +\ninformation'),(335,135,210,100,'Prediction'),(335,350,210,100,'Valuation'),(650,255,180,100,'Choice'),(920,255,180,100,'Outcome')]
    for x,y,w,h,l in coords:
        parts=l.split('\\n'); box(d,x,y,w,h,parts[0],parts[1] if len(parts)>1 else '',fill=COL['light'])
    for a in [(265,305,335,185),(265,305,335,400),(545,185,650,285),(545,400,650,325),(830,305,920,305)]: arrow(d,*a)
    box(d,650,500,240,90,'Feedback + learning','updates beliefs, values, options',fill=COL['sand'],stroke=COL['orange'])
    arrow(d,1010,355,870,500,stroke=COL['orange']); arrow(d,650,545,260,355,stroke=COL['orange'])
    txt(d,600,625,'Functional map, not a fixed conscious sequence',20,fill=COL['gray']); save_svg(d,p)
    figs['decision-loop']={'file':p.name,'caption':'Figure 1. The decision loop separates prediction from valuation and treats feedback as an input to future choices.','alt':'Flowchart from alternatives and information to prediction and valuation, then choice, outcome, and feedback.'}
    # 2 rational benchmark
    p=figdir/'rational-benchmark.svg'; d=new_svg(p,title='Rational-choice benchmark',desc='Feasible alternatives, beliefs and probabilities, and preferences and values combine in expected utility to guide choice.')
    txt(d,600,55,'Rational-choice benchmark under uncertainty',34,'bold')
    box(d,55,115,285,110,'Feasible alternatives','What can be done?'); box(d,55,275,285,110,'Beliefs / probabilities','What may follow?'); box(d,55,435,285,110,'Preferences / values','How good are outcomes?',fill=COL['sand'],stroke=COL['orange'])
    box(d,450,230,300,170,'Expected utility','compare probability-weighted outcomes',fill=COL['navy'],stroke=COL['navy']);
    # make text white overlay
    txt(d,600,300,'Expected utility',30,'bold',fill='white'); txt(d,600,340,'compare probability-weighted outcomes',17,fill='white')
    box(d,880,260,240,110,'Choice','best feasible option')
    for y in [170,330,490]: arrow(d,340,y,450,315)
    arrow(d,750,315,880,315)
    txt(d,600,610,'Benchmark, not biography: it makes assumptions and trade-offs explicit.',21,fill=COL['gray']); save_svg(d,p)
    figs['rational-benchmark']={'file':p.name,'caption':'Figure 2. Rational choice combines feasible actions, beliefs, and values; the diagram is a benchmark, not a literal mental algorithm.','alt':'Three inputs—feasible alternatives, beliefs and probabilities, preferences and values—feed expected utility and then choice.'}
    # 3 fast slow
    p=figdir/'fast-slow.svg'; d=new_svg(p,title='Fast and reflective processing',desc='Automatic processing generates impressions; reflective processing can inspect, redirect, and train it.')
    txt(d,600,55,'Fast answers and slow inspection',34,'bold')
    box(d,80,150,400,260,'Automatic processing','fast • associative • practiced',fill=COL['light'])
    txt(d,280,255,'impressions',26,'bold'); txt(d,280,300,'pattern recognition',23); txt(d,280,345,'action tendencies',23)
    box(d,720,150,400,260,'Reflective processing','effortful • rule-guided • limited',fill=COL['sand'],stroke=COL['orange'])
    txt(d,920,255,'inspection',26,'bold'); txt(d,920,300,'simulation + comparison',23); txt(d,920,345,'inhibition + redesign',23)
    arrow(d,480,250,720,250); arrow(d,720,340,480,340,stroke=COL['orange'])
    txt(d,600,500,'Practice can make slow skills fast; reflection can redirect attention and train intuition.',23)
    txt(d,600,555,'Neither mode is always rational. Reliability depends on the task, environment, and feedback.',20,fill=COL['gray']); save_svg(d,p)
    figs['fast-slow']={'file':p.name,'caption':'Figure 3. Automatic and reflective processing cooperate: one supplies candidate answers; the other can inspect, revise, and train them.','alt':'Two large boxes for automatic and reflective processing connected by arrows in both directions.'}
    # 4 attention filter
    p=figdir/'attention-filter.svg'; d=new_svg(p,title='Attention as gatekeeper',desc='A wide stream of world information passes through a narrow attention filter before becoming evidence for judgment.')
    txt(d,600,55,'Attention is the gatekeeper of evidence',34,'bold')
    for i,label in enumerate(['metrics','voices','threats','signals','memories','options','social cues']):
        txt(d,105,135+i*62,label,20,anchor='start',fill=COL['gray']); arrow(d,220,128+i*62,450,300,width=2)
    d.add(d.polygon(points=[(450,120),(650,250),(650,350),(450,520)],fill=COL['light'],stroke=COL['blue'],stroke_width=3))
    txt(d,535,292,'attention',26,'bold'); txt(d,535,326,'goals + salience',18,fill=COL['gray'])
    arrow(d,650,300,830,300)
    box(d,830,225,290,150,'Evidence for judgment','noticed • trusted • weighted',fill=COL['sand'],stroke=COL['orange'])
    txt(d,600,610,'What remains outside the filter may never enter the decision.',22,fill=COL['gray']); save_svg(d,p)
    figs['attention-filter']={'file':p.name,'caption':'Figure 4. Attention does not merely inspect evidence; it determines which information becomes evidence at all.','alt':'Many information streams narrow through an attention funnel into evidence for judgment.'}
    # 5 predictive loop
    p=figdir/'predictive-loop.svg'; d=new_svg(p,title='Predictive inference',desc='An internal model generates predictions; sensory evidence produces prediction error; precision weights the balance; action and learning update the model.')
    txt(d,600,48,'Predictive inference and active learning',34,'bold')
    box(d,80,210,230,125,'Internal model','priors • goals • causal beliefs'); box(d,485,95,230,120,'Predictions','expected sensations + outcomes'); box(d,485,390,230,120,'Sensory evidence','observations + feedback',fill=COL['sand'],stroke=COL['orange']); box(d,890,210,230,125,'Perception + action','best current inference')
    arrow(d,310,250,485,155); arrow(d,310,295,485,450,stroke=COL['orange']); arrow(d,715,155,890,250); arrow(d,715,450,890,295,stroke=COL['orange'])
    box(d,475,250,250,90,'Prediction error','mismatch weighted by precision',fill='white',stroke=COL['teal']); arrow(d,600,250,600,215,stroke=COL['teal']); arrow(d,600,390,600,340,stroke=COL['teal']); arrow(d,890,335,310,335,stroke=COL['teal'])
    txt(d,600,595,'Repeated, credible mismatch changes the model; action can also change the evidence received.',21,fill=COL['gray']); save_svg(d,p)
    figs['predictive-loop']={'file':p.name,'caption':'Figure 5. Perception and judgment emerge from predictions constrained by evidence; precision determines how strongly each source is weighted.','alt':'A loop linking internal model, predictions, sensory evidence, prediction error, perception and action, and model updating.'}
    # 6 valuation integration
    p=figdir/'valuation.svg'; d=new_svg(p,title='Constructed valuation',desc='Attributes, body state, goals, memory, identity, and social meaning feed a context-sensitive subjective value that guides approach and avoidance.')
    txt(d,600,48,'Value is constructed in context',34,'bold')
    inputs=[('attributes',110,140),('body state',110,250),('goals',110,360),('memory',110,470),('identity',940,140),('social meaning',940,250),('attention',940,360),('comparison',940,470)]
    for lab,x,y in inputs:
        box(d,x,y,160,70,lab,'',fill=COL['light'] if x<500 else COL['sand'],stroke=COL['blue'] if x<500 else COL['orange']);
        arrow(d,x+160 if x<500 else x, y+35, 500 if x<500 else 700,305,stroke=COL['blue'] if x<500 else COL['orange'],width=3)
    d.add(d.circle(center=(600,305),r=105,fill=COL['navy']))
    txt(d,600,295,'Subjective',27,'bold',fill='white'); txt(d,600,335,'value now',27,'bold',fill='white')
    arrow(d,600,410,600,540,stroke=COL['teal']); txt(d,600,585,'attention • motivation • choice',24,'bold',fill=COL['teal']); save_svg(d,p)
    figs['valuation']={'file':p.name,'caption':'Figure 6. Subjective value integrates option features with bodily state, goals, memory, identity, social meaning, attention, and comparison.','alt':'Multiple personal and contextual inputs converge on subjective value, which then guides attention, motivation, and choice.'}
    # 7 expectation loop
    p=figdir/'expectation-loop.svg'; d=new_svg(p,title='Self-fulfilling expectation loop',desc='Belief shapes attention, interpretation, emotion, action, other people and environment, feedback, and future belief.')
    txt(d,600,48,'When a prediction becomes part of the cause',34,'bold')
    labs=['Expectation','Attention','Interpretation','Emotion','Action','Environment / others','Feedback']
    cx,cy=600,330; R=230
    pts=[]
    for i,l in enumerate(labs):
        a=-math.pi/2+2*math.pi*i/len(labs); x=cx+R*math.cos(a); y=cy+R*math.sin(a); pts.append((x,y)); box(d,x-90,y-40,180,80,l,'',fill=COL['light'] if i<4 else COL['sand'],stroke=COL['blue'] if i<4 else COL['orange'])
    for i in range(len(pts)): arrow(d,pts[i][0]+(80 if i in [0,1,2] else 0),pts[i][1],pts[(i+1)%len(pts)][0],pts[(i+1)%len(pts)][1],stroke=COL['teal'],width=3)
    txt(d,600,330,'belief can',25,'bold'); txt(d,600,362,'change the system',25,'bold')
    save_svg(d,p)
    figs['expectation-loop']={'file':p.name,'caption':'Figure 7. Expectations can shape attention, action, and social response, producing feedback that appears to confirm the original belief.','alt':'Circular loop from expectation through attention, interpretation, emotion, action, environment and others, feedback, and back to expectation.'}
    # 8 heuristic substitution
    p=figdir/'heuristic-substitution.svg'; d=new_svg(p,title='Heuristic substitution',desc='A hard target question is silently replaced by an easier cue-based question, producing a quick answer that may or may not fit the target.')
    txt(d,600,55,'Heuristic substitution',34,'bold')
    box(d,60,180,330,160,'Target question','How likely is this project to succeed?',fill=COL['sand'],stroke=COL['orange']); arrow(d,390,260,535,260)
    box(d,535,155,330,210,'Easier question','How impressive does the founder feel?',fill=COL['light']); arrow(d,865,260,1030,260)
    box(d,1030,195,120,130,'Answer','“Likely”',fill=COL['navy'],stroke=COL['navy']); txt(d,1090,275,'Likely',24,'bold',fill='white')
    txt(d,600,470,'The shortcut is useful when the cue predicts the target—and biased when it does not.',24)
    txt(d,600,535,'Affect • availability • representativeness • recognition',21,fill=COL['gray']); save_svg(d,p)
    figs['heuristic-substitution']={'file':p.name,'caption':'Figure 8. A heuristic often answers an easier question than the one the decision actually requires.','alt':'A target question is replaced by an easier question and then produces a quick answer.'}
    # 9 natural frequencies
    p=figdir/'natural-frequencies.svg'; d=new_svg(p,title='Natural frequency tree',desc='Out of 10,000 people, 100 have a disease. Ninety test positive. Of 9,900 without it, 891 test positive. Therefore 90 of 981 positive tests are true positives, about 9 percent.')
    txt(d,600,48,'Why a “90% accurate” test can imply only 9% after a positive',32,'bold')
    box(d,40,245,180,95,'10,000 people','base population'); box(d,330,115,200,95,'100 diseased','1% base rate',fill=COL['sand'],stroke=COL['orange']); box(d,330,390,200,95,'9,900 not diseased','99%')
    arrow(d,220,285,330,160); arrow(d,220,305,330,435)
    box(d,690,85,210,85,'90 true positives','90% sensitivity',fill='#e7f4ea',stroke=COL['green']); box(d,690,205,210,85,'10 false negatives',''); box(d,690,360,210,85,'891 false positives','9% false-positive rate',fill='#f9e7e7',stroke=COL['red']); box(d,690,480,210,85,'9,009 true negatives','')
    arrow(d,530,160,690,125); arrow(d,530,165,690,245); arrow(d,530,435,690,400); arrow(d,530,445,690,520)
    box(d,980,225,180,170,'Positive tests','90 + 891 = 981',fill=COL['navy'],stroke=COL['navy']); txt(d,1070,302,'Positive tests',22,'bold',fill='white'); txt(d,1070,342,'90 + 891 = 981',18,fill='white')
    arrow(d,900,125,980,265,stroke=COL['green']); arrow(d,900,400,980,350,stroke=COL['red']); txt(d,1070,455,'P(disease | positive)',20,'bold'); txt(d,1070,493,'90 / 981 ≈ 9%',28,'bold',fill=COL['orange']); save_svg(d,p)
    figs['natural-frequencies']={'file':p.name,'caption':'Figure 9. Natural frequencies make base rates and false positives visible: only about 9% of positive tests in this example are true positives.','alt':'Frequency tree for 10,000 people showing 90 true positives and 891 false positives, producing a 9 percent posterior probability.'}
    # 10 context mechanisms
    p=figdir/'context-mechanisms.svg'; d=new_svg(p,title='Framing priming fluency and defaults',desc='Four context mechanisms: framing changes meaning; priming changes what is active; fluency changes ease; defaults change what happens without action.')
    txt(d,600,48,'Four ways context enters a choice',34,'bold')
    cards=[('Framing','What is the decision about?','meaning + reference point',70,130,COL['light'],COL['blue']),('Priming','What is already active?','associations + accessibility',620,130,COL['sand'],COL['orange']),('Fluency','How easy does it feel?','familiarity + confidence',70,365,'#edf4ed',COL['green']),('Default','What happens if I do nothing?','effort + implied norm',620,365,'#f5ebf5','#795a8a')]
    for title,q,sub,x,y,fill,stroke in cards:
        box(d,x,y,510,175,title,q,fill=fill,stroke=stroke); txt(d,x+255,y+135,sub,20,fill=COL['gray'])
    save_svg(d,p)
    figs['context-mechanisms']={'file':p.name,'caption':'Figure 10. Framing, priming, fluency, and defaults influence choice through different mechanisms and should not be treated as synonyms.','alt':'Four cards distinguish framing, priming, fluency, and defaults by the question each mechanism answers.'}
    # 11 narrator-learning
    p=figdir/'narrator-learning.svg'; d=new_svg(p,title='Choice, explanation, and learning',desc='A choice produces an outcome; the narrator explains it. Hindsight and outcome bias can distort learning unless a contemporaneous decision record is compared with results.')
    txt(d,600,48,'The narrator can protect identity—or protect learning',34,'bold')
    box(d,50,220,180,100,'Choice','reasons felt now'); arrow(d,230,270,390,270); box(d,390,220,180,100,'Outcome','luck + process'); arrow(d,570,270,730,270); box(d,730,220,200,100,'Narrative','“why it happened”'); arrow(d,930,270,1090,270); box(d,1010,400,160,100,'Future rule','what repeats')
    box(d,390,430,300,105,'Decision record','beliefs • probabilities • values',fill=COL['sand'],stroke=COL['orange']); arrow(d,540,430,810,320,stroke=COL['orange']); arrow(d,690,482,1010,450,stroke=COL['orange'])
    txt(d,600,595,'A record written before the outcome prevents luck from rewriting the process.',22,fill=COL['gray']); save_svg(d,p)
    figs['narrator-learning']={'file':p.name,'caption':'Figure 11. Learning improves when post-outcome explanations are checked against a decision record created before hindsight arrived.','alt':'Choice leads to outcome, narrative, and future rule, while a decision record provides an independent comparison.'}
    # 12 habit loop
    p=figdir/'habit-loop.svg'; d=new_svg(p,title='Habit loop',desc='Cue triggers urge, response, reward or relief, and learning, which strengthens the cue-response association.')
    txt(d,600,48,'A habit is a context-bound decision loop',34,'bold')
    labs=[('Cue','time • place • emotion'),('Urge','wanting • tension'),('Response','what the body does'),('Reward / relief','pleasure • escape • completion'),('Learning','cue predicts response')]
    xs=[55,285,525,765,1000]
    for i,(a,b) in enumerate(labs): box(d,xs[i],245,175 if i<4 else 150,115,a,b,fill=COL['light'] if i<3 else COL['sand'],stroke=COL['blue'] if i<3 else COL['orange']);
    for i in range(4): arrow(d,xs[i]+(175 if i<4 else 150),302,xs[i+1],302)
    arrow(d,1075,360,150,470,stroke=COL['teal']); arrow(d,150,470,150,360,stroke=COL['teal'])
    txt(d,600,485,'Repetition in a stable context compresses prediction and valuation into a cached response.',22)
    txt(d,600,545,'Relief can reinforce behavior even when long-term outcomes worsen.',20,fill=COL['gray']); save_svg(d,p)
    figs['habit-loop']={'file':p.name,'caption':'Figure 12. Habits strengthen when cues repeatedly trigger responses that deliver reward or relief.','alt':'Linear habit loop from cue to urge, response, reward or relief, and learning, with a feedback arrow back to cue.'}
    # 13 behavior design
    p=figdir/'behavior-design.svg'; d=new_svg(p,title='Behavior design stack',desc='Behavior is supported by motivation, ability and prompt, with upstream tools: environment, friction, plans, social support, and recovery.')
    txt(d,600,48,'Design behavior upstream of the moment of temptation',34,'bold')
    layers=[('5. Recovery','relapse plan • restart at next cue',120,500,COL['light'],COL['blue']),('4. Prompt','visible cue at the right moment',180,405,COL['light'],COL['blue']),('3. Ability','reduce friction • smallest viable action',240,310,'#edf4ed',COL['green']),('2. Motivation','meaning • immediate reward • commitment',300,215,COL['sand'],COL['orange']),('1. Environment','availability • defaults • social context',360,120,'#f5ebf5','#795a8a')]
    for title,sub,x,y,fill,stroke in layers: box(d,x,y,1200-2*x,75,title,sub,fill=fill,stroke=stroke)
    txt(d,600,610,'The most reliable intervention is often the one that requires the least heroism.',22,fill=COL['gray']); save_svg(d,p)
    figs['behavior-design']={'file':p.name,'caption':'Figure 13. Behavior design works best upstream: environment, motivation, ability, prompts, and recovery plans reinforce one another.','alt':'A five-layer stack from environment through motivation, ability, prompt, and recovery.'}
    # 14 social pathways
    p=figdir/'social-pathways.svg'; d=new_svg(p,title='Social influence pathways',desc='Other people influence judgment through information, norms and belonging, authority, reciprocity and commitment, and identity and similarity.')
    txt(d,600,48,'How other minds enter the decision loop',34,'bold')
    d.add(d.circle(center=(600,330),r=105,fill=COL['navy'])); txt(d,600,320,'Choice',30,'bold',fill='white'); txt(d,600,360,'and action',25,'bold',fill='white')
    nodes=[('Information','What do they know?',170,120),('Norms','What do people do?',170,455),('Authority','Who has legitimate expertise?',860,100),('Reciprocity + commitment','What do I owe or stand for?',815,455),('Identity + similarity','What do people like us do?',490,535)]
    for lab,sub,x,y in nodes:
        box(d,x,y,260 if x!=490 else 300,90,lab,sub,fill=COL['light'] if x<500 else COL['sand'],stroke=COL['blue'] if x<500 else COL['orange']); arrow(d,x+(260 if x!=490 else 150),y+45,600,330,width=3)
    save_svg(d,p)
    figs['social-pathways']={'file':p.name,'caption':'Figure 14. Social influence enters through information, norms, authority, reciprocity, commitment, identity, and similarity.','alt':'Five social influence pathways point toward choice and action.'}
    # 15 persuasion update
    p=figdir/'persuasion-update.svg'; d=new_svg(p,title='Persuasion as model updating',desc='An audience current model meets a credible message that directs attention, provides evidence, connects value, addresses identity and efficacy, and makes action possible.')
    txt(d,600,48,'Persuasion is guided model updating',34,'bold')
    box(d,45,230,230,120,'Audience model','beliefs • values • identity'); arrow(d,275,290,430,290)
    box(d,430,145,340,290,'Message design','ethos • pathos • logos • kairos',fill=COL['light'])
    for i,s in enumerate(['attention','evidence','meaning','efficacy']): txt(d,600,245+i*42,'• '+s,21,anchor='middle')
    arrow(d,770,290,925,290); box(d,925,230,230,120,'Updated model','belief + intention + action',fill=COL['sand'],stroke=COL['orange'])
    txt(d,600,520,'A message fails when it answers the obstacle the persuader wishes existed.',22)
    txt(d,600,565,'Ethical persuasion preserves the audience’s ability to inspect and refuse.',20,fill=COL['gray']); save_svg(d,p)
    figs['persuasion-update']={'file':p.name,'caption':'Figure 15. Persuasion succeeds when message design addresses the audience’s actual model and obstacle to change.','alt':'Audience model flows through message design—attention, evidence, meaning and efficacy—to an updated model and action.'}
    # 16 story update
    p=figdir/'story-update.svg'; d=new_svg(p,title='Story as guided model updating',desc='Audience current story moves through trouble or prediction error, protagonist conflict, new model, evidence and action.')
    txt(d,600,48,'Story is model change with a protagonist',34,'bold')
    labels=[('Current story','what the audience expects'),('BUT','trouble / prediction error'),('Conflict','a protagonist chooses'),('New model','what the old story missed'),('THEREFORE','evidence + action')]
    x=35
    widths=[205,205,205,205,205]
    for i,(a,b) in enumerate(labels):
        fill=COL['light'] if i<3 else COL['sand']; stroke=COL['blue'] if i<3 else COL['orange']; box(d,x,230,widths[i],135,a,b,fill=fill,stroke=stroke)
        if i<4: arrow(d,x+widths[i],298,x+widths[i]+30,298)
        x+=widths[i]+30
    txt(d,600,485,'AND establishes shared reality. BUT creates meaningful mismatch. THEREFORE offers a credible path.',22)
    txt(d,600,540,'Evidence establishes scale and reliability; story creates meaning and simulation.',20,fill=COL['gray']); save_svg(d,p)
    figs['story-update']={'file':p.name,'caption':'Figure 16. A persuasive story begins in the audience’s current model, creates meaningful trouble, and resolves in an evidence-supported new model and action.','alt':'Five-step story sequence: current story, but or trouble, protagonist conflict, new model, therefore evidence and action.'}
    # 17 communication grounding
    p=figdir/'communication-grounding.svg'; d=new_svg(p,title='Communication as grounding',desc='Speaker expresses a model; listener interprets it; questions, paraphrase and repair build common ground and shared action.')
    txt(d,600,48,'Communication is joint inference, not message delivery',34,'bold')
    box(d,55,180,260,140,'Person A','intention + context'); box(d,885,180,260,140,'Person B','interpretation + context',fill=COL['sand'],stroke=COL['orange'])
    box(d,450,130,300,110,'Words + signals','content • tone • timing'); arrow(d,315,225,450,190); arrow(d,750,190,885,225)
    box(d,450,370,300,120,'Grounding + repair','ask • paraphrase • correct',fill='#edf4ed',stroke=COL['green']); arrow(d,885,305,750,420,stroke=COL['green']); arrow(d,450,420,315,305,stroke=COL['green'])
    txt(d,600,560,'Understanding is good enough for the current purpose—not assumed once and for all.',21,fill=COL['gray']); save_svg(d,p)
    figs['communication-grounding']={'file':p.name,'caption':'Figure 17. Human communication requires grounding: participants check, correct, and repair interpretations until common ground is sufficient for action.','alt':'Two people exchange words and signals, then use grounding and repair in a feedback loop.'}
    # 18 negotiation architecture
    p=figdir/'negotiation-architecture.svg'; d=new_svg(p,title='Negotiation architecture',desc='Each side has BATNA, reservation value, interests and priorities. Communication and offers search a possible agreement space where value can be claimed and created.')
    txt(d,600,48,'Negotiation is joint decision-making',34,'bold')
    box(d,40,135,300,340,'Party A','BATNA • reservation • interests',fill=COL['light']); box(d,860,135,300,340,'Party B','BATNA • reservation • interests',fill=COL['sand'],stroke=COL['orange'])
    txt(d,190,270,'Alternatives',23,'bold'); txt(d,190,315,'Priorities',23); txt(d,190,360,'Constraints',23)
    txt(d,1010,270,'Alternatives',23,'bold'); txt(d,1010,315,'Priorities',23); txt(d,1010,360,'Constraints',23)
    box(d,440,135,320,150,'Information + process','questions • standards • trust',fill='#edf4ed',stroke=COL['green']); box(d,440,360,320,150,'Agreement space','claim value + create value',fill=COL['navy'],stroke=COL['navy']); txt(d,600,425,'Agreement space',27,'bold',fill='white'); txt(d,600,465,'claim + create value',19,fill='white')
    arrow(d,340,250,440,210); arrow(d,860,250,760,210,stroke=COL['orange']); arrow(d,600,285,600,360,stroke=COL['green']); arrow(d,440,440,340,390,stroke=COL['green']); arrow(d,760,440,860,390,stroke=COL['green'])
    txt(d,600,585,'No agreement is better than an agreement worse than a credible alternative.',22,fill=COL['gray']); save_svg(d,p)
    figs['negotiation-architecture']={'file':p.name,'caption':'Figure 18. Negotiation links each party’s alternatives and priorities to a jointly designed agreement space.','alt':'Party A and Party B feed information and process into an agreement space where value is claimed and created.'}
    # 19 ZOPA
    p=figdir/'zopa.svg'; d=new_svg(p,title='Zone of possible agreement',desc='Seller reservation value at 8,000 euros and buyer reservation value at 10,000 euros create a 2,000-euro zone of possible agreement.')
    txt(d,600,55,'The zone of possible agreement (ZOPA)',34,'bold')
    y=300; d.add(d.line(start=(100,y),end=(1100,y),stroke=COL['navy'],stroke_width=8))
    for val,x in [(7000,180),(8000,380),(9000,600),(10000,820),(11000,1020)]:
        d.add(d.line(start=(x,y-18),end=(x,y+18),stroke=COL['navy'],stroke_width=4)); txt(d,x,y+58,f'€{val:,}',20)
    d.add(d.rect(insert=(380,y-38),size=(440,76),fill=COL['light'],stroke=COL['blue'],stroke_width=2,opacity=.8))
    txt(d,600,y-70,'Bargaining surplus = €2,000',24,'bold')
    txt(d,380,y+115,'Seller reservation',20,'bold',fill=COL['orange']); txt(d,820,y+115,'Buyer reservation',20,'bold',fill=COL['blue'])
    txt(d,600,520,'A negative ZOPA is not a failure of tactics. It is a signal to walk away or add issues.',22,fill=COL['gray']); save_svg(d,p)
    figs['zopa']={'file':p.name,'caption':'Figure 19. A positive bargaining zone exists when the buyer’s maximum exceeds the seller’s minimum; the overlap is the bargaining surplus.','alt':'Number line showing seller reservation at 8,000 euros, buyer reservation at 10,000 euros, and a 2,000-euro zone between them.'}
    # 20 integrative frontier
    p=figdir/'pareto.svg'; d=new_svg(p,title='Pareto frontier',desc='Negotiated agreements inside a curved frontier are inefficient; moving northeast can improve both parties until reaching the Pareto frontier. Distribution along the frontier remains a fairness question.')
    txt(d,600,48,'Create value before arguing over its division',34,'bold')
    # axes
    d.add(d.line(start=(150,540),end=(1080,540),stroke=COL['navy'],stroke_width=4)); d.add(d.line(start=(150,540),end=(150,100),stroke=COL['navy'],stroke_width=4)); arrow(d,1080,540,1120,540); arrow(d,150,100,150,70)
    txt(d,650,610,'Value to Party A',22,'bold'); txt(d,55,310,'Value to Party B',22,'bold',anchor='middle')
    # frontier curve
    path=d.path(d='M 250 500 C 420 420, 650 260, 990 140',fill='none',stroke=COL['orange'],stroke_width=8); d.add(path)
    txt(d,820,165,'Pareto frontier',24,'bold',fill=COL['orange'])
    # points and arrows
    for x,y,l in [(390,430,'compromise'),(500,360,'better package'),(760,245,'efficient deal')]:
        d.add(d.circle(center=(x,y),r=11,fill=COL['blue'])); txt(d,x+18,y-15,l,18,anchor='start')
    arrow(d,390,430,500,360,stroke=COL['green']); arrow(d,500,360,760,245,stroke=COL['green'])
    txt(d,600,565,'Efficiency asks whether value is wasted. Fairness asks how the created value is divided.',20,fill=COL['gray']); save_svg(d,p)
    figs['pareto']={'file':p.name,'caption':'Figure 20. Integrative negotiation moves agreements toward the Pareto frontier; efficiency and fairness remain distinct questions.','alt':'Graph with value to two parties on the axes and a Pareto frontier; arrows move from compromise to better package to efficient deal.'}
    # 21 DPN cycle
    p=figdir/'dpn-cycle.svg'; d=new_svg(p,title='DPN cycle',desc='Notice, test, ask, design and learn form a repeating cycle for decision, persuasion and negotiation.')
    txt(d,600,48,'The DPN cycle',36,'bold')
    labs=[('NOTICE','What entered attention?'),('TEST','What evidence could change the model?'),('ASK','Whose perspective or interest is missing?'),('DESIGN','What process or environment helps?'),('LEARN','What did the result reveal?')]
    cx,cy,R=600,330,230; pts=[]
    for i,(a,b) in enumerate(labs):
        ang=-math.pi/2+2*math.pi*i/5; x=cx+R*math.cos(ang); y=cy+R*math.sin(ang); pts.append((x,y)); box(d,x-105,y-47,210,94,a,b,fill=COL['light'] if i<3 else COL['sand'],stroke=COL['blue'] if i<3 else COL['orange'])
    for i in range(5): arrow(d,pts[i][0],pts[i][1],pts[(i+1)%5][0],pts[(i+1)%5][1],stroke=COL['teal'],width=3)
    d.add(d.circle(center=(cx,cy),r=85,fill=COL['navy'])); txt(d,cx,cy-4,'Decision',24,'bold',fill='white'); txt(d,cx,cy+26,'hygiene',24,'bold',fill='white')
    save_svg(d,p)
    figs['dpn-cycle']={'file':p.name,'caption':'Figure 21. Notice, test, ask, design, and learn form a repeatable cycle for improving decisions, persuasion, communication, and negotiation.','alt':'Circular five-step cycle: notice, test, ask, design, and learn, around decision hygiene.'}
    return figs

FIGURE_PLACEMENT={
 1:('The decision loop','decision-loop'), 2:('A family of models','rational-benchmark'), 4:('Fast thinking is not primitive thinking','fast-slow'),
 5:('Every organism inhabits a partial world','attention-filter'), 6:('Prediction, error, and precision','predictive-loop'), 7:('Emotion tells the mind what matters','valuation'),
 8:('Expectation in the body','expectation-loop'), 9:('Trigger features and cue-based judgment','heuristic-substitution'), 12:('Conditional probability and the base-rate trap','natural-frequencies'),
 15:('Equivalent facts, different meanings','context-mechanisms'), 18:('Making the narrator accountable','narrator-learning'), 19:('How habits form','habit-loop'),
 21:('Choosing the right tool','behavior-design'), 23:('Social proof and cascades','social-pathways'), 25:('Two routes through a message','persuasion-update'),
 26:('Story as model updating','story-update'), 28:('Communication is coordination and repair','communication-grounding'), 30:('What negotiation is not','negotiation-architecture'),
 31:('The structure of distributive bargaining','zopa'), 33:('Pareto efficiency','pareto'), 35:('The DPN cycle: notice, test, ask, design, learn','dpn-cycle')
}

# ---------- markdown ----------

def inline_md(p: PBlock) -> str:
    s=p.text.strip()
    if not s: return ''
    if p.bold: return f'**{s}**'
    if p.italic: return f'*{s}*'
    return s


def _html_breaks(s: str) -> str:
    return re.sub(r'\s*<br\s*/?>\s*','<br />',s,flags=re.I)


def _callout_html(label: str, body: str) -> str:
    kind=slugify(label)
    body_html=_html_breaks(html.escape(body).replace('&lt;br&gt;','<br />').replace('&lt;br /&gt;','<br />'))
    return f'<aside class="callout {html.escape(kind)}"><div class="callout-title">{html.escape(label)}</div><div class="callout-body">{body_html}</div></aside>'


def table_to_md(t: TBlock) -> str:
    rows=t.rows
    if not rows: return ''
    maxc=max(len(r) for r in rows)
    if maxc==1:
        vals=[r[0] if r else '' for r in rows]
        # Some callouts arrived as one cell with an explicit line break.
        if len(vals)==1 and re.search(r'<br\s*/?>',vals[0],re.I):
            label,body=re.split(r'<br\s*/?>',vals[0],maxsplit=1,flags=re.I)
            if label.strip() and (label.strip().upper()==label.strip() or label.strip().lower() in {'core idea','scientific caution','activity','evidence and boundary conditions'}):
                return _callout_html(label.strip(),body.strip())
        if len(vals)>=2 and vals[0].strip():
            return _callout_html(vals[0].strip(),'\n\n'.join(vals[1:]))
        return _html_breaks('\n\n'.join(vals))
    rows=[r+['']*(maxc-len(r)) for r in rows]
    header=rows[0]
    def cell(c): return _html_breaks(esc_md(c))
    lines=['| '+' | '.join(cell(c) for c in header)+' |','| '+' | '.join('---' for _ in header)+' |']
    for row in rows[1:]: lines.append('| '+' | '.join(cell(c) for c in row)+' |')
    return '\n'.join(lines)


def blocks_to_md(ch: Chapter, figs: dict, for_epub=False, include_refs=True, image_prefix='../figures/') -> str:
    lines=[]
    lines.append('---')
    lines.append(f'title: "{ch.title.replace(chr(34), chr(39))}"')
    lines.append(f'subtitle: "{ch.subtitle.replace(chr(34), chr(39))}"')
    lines.append(f'chapter: {ch.num}')
    lines.append(f'part: {ch.part}')
    lines.append(f'slug: "{ch.slug}"')
    lines.append('---\n')
    lines.append(f'# Chapter {ch.num}. {ch.title}\n')
    lines.append(f'*{ch.subtitle}*\n')
    placement=FIGURE_PLACEMENT.get(ch.num)
    inserted=False
    in_list=False
    def close_list():
        nonlocal in_list
        if in_list:
            lines.append('')
            in_list=False
    def figure_html(f):
        ext='png' if for_epub else 'svg'
        name=Path(f['file']).with_suffix('.'+ext).name
        return f'<figure class="book-figure"><img src="{image_prefix}{name}" alt="{html.escape(f["alt"],quote=True)}" /><figcaption>{html.escape(f["caption"])}</figcaption></figure>\n'
    for b in ch.blocks:
        if isinstance(b,PBlock):
            t=b.text.strip(); st=b.style
            if not t: continue
            if st in {'Learning Objective','Takeaway Bullet'}:
                item=re.sub(r'^[•\-–—]\s*','',t)
                lines.append(f'- {item}')
                in_list=True
                continue
            close_list()
            if st=='Heading 2':
                lines.append(f'## {t}\n')
                if placement and placement[0].lower() in t.lower() and not inserted:
                    lines.append(figure_html(figs[placement[1]])); inserted=True
            elif st=='Heading 3': lines.append(f'### {t}\n')
            elif st=='Caption': lines.append(f'*{t}*\n')
            elif st in {'Chapter Label','Chapter Title','Chapter Subtitle','Part Title','TOC Entry'}: pass
            else: lines.append(inline_md(b)+'\n')
        else:
            close_list(); lines.append(table_to_md(b)+'\n')
    close_list()
    if placement and not inserted:
        lines.insert(10,figure_html(figs[placement[1]]))
    if include_refs and ch.references:
        lines.append('## References cited in this chapter\n')
        for ref in ch.references: lines.append(f'<div class="reference">{html.escape(ref)}</div>\n')
    return '\n'.join(lines).strip()+"\n"


def appendix_to_md(title: str, blocks: List[Block]) -> str:
    lines=['---',f'title: "{title}"','appendix: true','---\n',f'# {title}\n']
    in_list=False
    for b in blocks:
        if isinstance(b,PBlock):
            if not b.text.strip(): continue
            if b.style in {'Learning Objective','Takeaway Bullet'}:
                item=re.sub(r'^[•\-–—]\s*','',b.text)
                lines.append(f'- {item}'); in_list=True; continue
            if in_list: lines.append(''); in_list=False
            if b.style=='Heading 2': lines.append(f'## {b.text}\n')
            elif b.style=='Heading 3': lines.append(f'### {b.text}\n')
            else: lines.append(b.text+'\n')
        else:
            if in_list: lines.append(''); in_list=False
            lines.append(table_to_md(b)+'\n')
    if in_list: lines.append('')
    return '\n'.join(lines)

# ---------- site ----------
SITE_TEMPLATE='''<!doctype html>
<html lang="en" data-theme="light">
<head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{{ page_title }} · Decision in the Making</title>
<meta name="description" content="{{ description|e }}">
<link rel="stylesheet" href="{{ root }}assets/style.css">
<script defer src="{{ root }}assets/site.js"></script>
</head>
<body>
<a class="skip-link" href="#main">Skip to content</a>
<header class="topbar"><button class="menu-button" aria-label="Open navigation">☰</button><a class="brand" href="{{ root }}index.html">Decision in the Making</a><div class="top-actions"><button id="search-button" aria-label="Search book">Search</button><button id="theme-button" aria-label="Toggle dark mode">◐</button></div></header>
<div class="layout">
<aside class="sidebar" aria-label="Book navigation"><div class="sidebar-inner"><a class="course-home" href="{{ root }}index.html">Course textbook</a>{{ nav|safe }}</div></aside>
<main id="main" class="content">{{ content|safe }}<nav class="page-nav">{% if prev %}<a class="prev" href="{{ root }}{{ prev.href }}">← {{ prev.label }}</a>{% endif %}{% if next %}<a class="next" href="{{ root }}{{ next.href }}">{{ next.label }} →</a>{% endif %}</nav></main>
<aside class="on-page" aria-label="On this page">{{ toc|safe }}</aside>
</div>
<dialog id="search-dialog"><form method="dialog"><button class="close-search">Close</button></form><h2>Search the textbook</h2><input id="search-input" type="search" placeholder="Search concepts, cases, and tools" autocomplete="off"><div id="search-results"></div></dialog>
<footer><p>© 2026 Huanren Warren Zhang. <em>Decision in the Making</em>.</p></footer>
</body></html>'''

CSS='''
:root{--ink:#183047;--muted:#607080;--paper:#fff;--panel:#f4f7fa;--line:#d7e0e7;--accent:#25678f;--accent2:#b95f2d;--max:780px;--side:292px;--toc:230px;color-scheme:light dark}
html[data-theme="dark"]{--ink:#e9f1f7;--muted:#aebdca;--paper:#111922;--panel:#18232e;--line:#304250;--accent:#78bde5;--accent2:#f3a16d;color-scheme:dark}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;background:var(--paper);color:var(--ink);font-family:Inter,ui-sans-serif,system-ui,-apple-system,"Segoe UI",sans-serif;line-height:1.72}.skip-link{position:absolute;left:-9999px}.skip-link:focus{left:1rem;top:1rem;z-index:99;background:var(--paper);padding:.6rem}.topbar{position:sticky;top:0;z-index:20;height:58px;display:flex;align-items:center;gap:1rem;border-bottom:1px solid var(--line);background:color-mix(in srgb,var(--paper) 94%,transparent);backdrop-filter:blur(12px);padding:0 1rem}.brand{font-weight:750;color:var(--ink);text-decoration:none}.menu-button{display:none}.top-actions{margin-left:auto;display:flex;gap:.5rem}button{font:inherit;border:1px solid var(--line);border-radius:8px;background:var(--panel);color:var(--ink);padding:.42rem .75rem;cursor:pointer}.layout{display:grid;grid-template-columns:var(--side) minmax(0,1fr) var(--toc);max-width:1500px;margin:auto}.sidebar{border-right:1px solid var(--line);min-height:calc(100vh - 58px)}.sidebar-inner{position:sticky;top:58px;max-height:calc(100vh - 58px);overflow:auto;padding:1.3rem 1rem 4rem}.course-home{display:block;font-weight:800;color:var(--accent);text-decoration:none;margin-bottom:1rem}.part-nav{margin:1rem 0}.part-nav>summary{font-weight:750;cursor:pointer}.part-nav a{display:block;color:var(--muted);text-decoration:none;padding:.3rem 0 .3rem .65rem;border-left:2px solid transparent;font-size:.92rem}.part-nav a:hover,.part-nav a.active{color:var(--accent);border-left-color:var(--accent)}.content{width:min(var(--max),calc(100% - 3rem));margin:0 auto;padding:3.4rem 0 6rem;min-width:0}.on-page{padding:2rem 1rem}.on-page nav{position:sticky;top:80px;border-left:1px solid var(--line);padding-left:1rem;font-size:.88rem}.on-page a{display:block;color:var(--muted);text-decoration:none;padding:.2rem 0}.on-page a:hover{color:var(--accent)}h1,h2,h3{line-height:1.22;scroll-margin-top:80px}h1{font-family:Georgia,serif;font-size:clamp(2.15rem,5vw,3.6rem);letter-spacing:-.025em;margin:0 0 .5rem}h2{font-size:1.65rem;margin-top:3rem;border-top:1px solid var(--line);padding-top:1.4rem}h3{font-size:1.18rem;margin-top:2rem}p{margin:1rem 0}a{color:var(--accent)}blockquote,.callout{border-left:5px solid var(--accent);background:var(--panel);padding:1rem 1.2rem;margin:1.5rem 0;border-radius:0 10px 10px 0}.callout-title{font-size:.8rem;text-transform:uppercase;letter-spacing:.08em;font-weight:850;color:var(--accent)}.callout-body{margin-top:.45rem}.core-idea{border-left-color:var(--accent2)}.core-idea .callout-title{color:var(--accent2)}table{display:block;width:100%;overflow-x:auto;border-collapse:collapse;margin:1.4rem 0;font-size:.93rem}th,td{border:1px solid var(--line);padding:.62rem .7rem;vertical-align:top}th{background:var(--panel);text-align:left}.book-figure{margin:2rem 0}.book-figure img{width:100%;height:auto;border:1px solid var(--line);border-radius:12px;background:white}.book-figure figcaption{font-size:.88rem;color:var(--muted);margin-top:.6rem}.reference{padding-left:1.5rem;text-indent:-1.5rem;margin:.55rem 0;font-size:.92rem}.page-nav{display:flex;justify-content:space-between;gap:1rem;border-top:1px solid var(--line);margin-top:4rem;padding-top:1.5rem}.page-nav a{text-decoration:none;font-weight:700}.next{text-align:right;margin-left:auto}.chapter-card-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:1rem}.chapter-card{display:block;text-decoration:none;color:var(--ink);border:1px solid var(--line);border-radius:12px;padding:1rem;background:var(--panel)}.chapter-card small{color:var(--accent);font-weight:750}.chapter-card strong{display:block;margin:.25rem 0}.hero{padding:2.5rem;border:1px solid var(--line);border-radius:18px;background:linear-gradient(135deg,var(--panel),var(--paper))}.hero p{font-size:1.1rem}.part-intro{margin-top:3rem}.download-row{display:flex;flex-wrap:wrap;gap:.7rem}.download-row a{border:1px solid var(--line);background:var(--panel);padding:.55rem .8rem;border-radius:8px;text-decoration:none;font-weight:700}dialog{width:min(760px,94vw);border:1px solid var(--line);border-radius:14px;background:var(--paper);color:var(--ink);padding:1.5rem}dialog::backdrop{background:rgba(0,0,0,.5)}#search-input{width:100%;font:inherit;padding:.8rem;border:1px solid var(--line);border-radius:8px;background:var(--paper);color:var(--ink)}.search-hit{padding:.8rem 0;border-bottom:1px solid var(--line)}.search-hit a{font-weight:750;text-decoration:none}.close-search{float:right}footer{border-top:1px solid var(--line);padding:1.5rem;text-align:center;color:var(--muted);font-size:.85rem}
@media(max-width:1100px){.layout{grid-template-columns:var(--side) minmax(0,1fr)}.on-page{display:none}}
@media(max-width:780px){.menu-button{display:inline-block}.layout{display:block}.sidebar{position:fixed;z-index:30;left:-310px;top:58px;width:300px;background:var(--paper);transition:left .2s}.sidebar.open{left:0;box-shadow:0 10px 35px rgba(0,0,0,.25)}.content{width:min(100% - 2rem,var(--max));padding-top:2rem}.brand{font-size:.9rem}.top-actions button:first-child{font-size:0}.top-actions button:first-child:after{content:'⌕';font-size:1rem}.hero{padding:1.3rem}}
@media print{.topbar,.sidebar,.on-page,.page-nav,dialog,footer{display:none!important}.layout{display:block}.content{width:100%;max-width:none;padding:0}a{color:inherit;text-decoration:none}h2{break-after:avoid}.book-figure,table,.callout{break-inside:avoid}}
'''

JS='''
const root=document.documentElement;const saved=localStorage.getItem('dpn-theme');if(saved)root.dataset.theme=saved;
document.getElementById('theme-button')?.addEventListener('click',()=>{root.dataset.theme=root.dataset.theme==='dark'?'light':'dark';localStorage.setItem('dpn-theme',root.dataset.theme)});
const sidebar=document.querySelector('.sidebar');document.querySelector('.menu-button')?.addEventListener('click',()=>sidebar.classList.toggle('open'));
const dlg=document.getElementById('search-dialog'),btn=document.getElementById('search-button'),inp=document.getElementById('search-input'),out=document.getElementById('search-results');
btn?.addEventListener('click',()=>{dlg.showModal();setTimeout(()=>inp.focus(),60)});
let idx=[];fetch((document.body.dataset.root||'')+'search-index.json').then(r=>r.json()).then(x=>idx=x).catch(()=>{});
function esc(s){return s.replace(/[&<>"']/g,m=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[m]))}
inp?.addEventListener('input',()=>{let q=inp.value.trim().toLowerCase();if(q.length<2){out.innerHTML='';return}let terms=q.split(/\s+/);let hits=idx.map(x=>{let hay=(x.title+' '+x.headings+' '+x.text).toLowerCase();let score=terms.reduce((s,t)=>s+(x.title.toLowerCase().includes(t)?5:0)+(x.headings.toLowerCase().includes(t)?3:0)+(hay.split(t).length-1),0);return {...x,score}}).filter(x=>x.score>0).sort((a,b)=>b.score-a.score).slice(0,25);out.innerHTML=hits.map(x=>`<div class="search-hit"><a href="${document.body.dataset.root||''}${x.href}">${esc(x.title)}</a><div>${esc(x.excerpt)}</div></div>`).join('')||'<p>No results.</p>'});
'''


def nav_html(chapters, appendices, active='', prefix=''):
    s=[]
    for part,(pt,_) in PARTS.items():
        s.append(f'<details class="part-nav" open><summary>Part {part}. {html.escape(pt)}</summary>')
        for ch in [x for x in chapters if x.part==part]:
            cls=' class="active"' if ch.slug==active else ''
            s.append(f'<a{cls} href="{prefix}chapters/{ch.slug}.html">{ch.num}. {html.escape(ch.title)}</a>')
        s.append('</details>')
    s.append('<details class="part-nav" open><summary>Appendices</summary>')
    for slug,title in appendices: s.append(f'<a href="{prefix}appendices/{slug}.html">{html.escape(title)}</a>')
    s.append(f'<a href="{prefix}references.html">References</a><a href="{prefix}about.html">About and editorial notes</a></details>')
    return ''.join(s)


def headings_toc(soup):
    hs=soup.find_all(['h2','h3'])
    if not hs: return ''
    arr=['<nav><strong>On this page</strong>']
    used={}
    for h in hs:
        sid=slugify(h.get_text(' ',strip=True)); n=used.get(sid,0)+1; used[sid]=n
        if n>1: sid=f'{sid}-{n}'
        h['id']=sid
        cls=' class="toc-h3"' if h.name=='h3' else ''
        arr.append(f'<a{cls} href="#{sid}">{html.escape(h.get_text(" ",strip=True))}</a>')
    arr.append('</nav>')
    return ''.join(arr)


def render_markdown(md: MarkdownIt, text: str):
    # Remove YAML front matter before markdown rendering.
    text=re.sub(r'^---\n.*?\n---\n','',text,flags=re.S)
    return md.render(text)


def build_site(chapters, appendices, all_refs, figs):
    docs=OUT/'docs'; assets=docs/'assets'; chdir=docs/'chapters'; apdir=docs/'appendices'; fdir=docs/'figures'; down=docs/'downloads'
    for p in [assets,chdir,apdir,fdir,down]: p.mkdir(parents=True,exist_ok=True)
    shutil.copytree(OUT/'figures',fdir,dirs_exist_ok=True)
    (assets/'style.css').write_text(CSS,encoding='utf-8'); (assets/'site.js').write_text(JS,encoding='utf-8'); (docs/'.nojekyll').write_text('',encoding='utf-8')
    md=MarkdownIt('commonmark',{'html':True,'linkify':True}).enable('table')
    env=Environment(loader=BaseLoader(),autoescape=select_autoescape(['html'])); tpl=env.from_string(SITE_TEMPLATE)
    appmeta=[(slugify(t),t) for t in appendices]
    nav=nav_html(chapters,appmeta)
    pages=[]
    for ch in chapters: pages.append({'href':f'chapters/{ch.slug}.html','label':f'Chapter {ch.num}. {ch.title}','ch':ch})
    for title in appendices: pages.append({'href':f'appendices/{slugify(title)}.html','label':title,'app':title})
    pages += [{'href':'references.html','label':'References'},{'href':'about.html','label':'About and editorial notes'}]
    search=[]
    # chapter pages
    for i,ch in enumerate(chapters):
        src=(OUT/'chapters'/f'{ch.slug}.md').read_text(encoding='utf-8')
        content=render_markdown(md,src); soup=BeautifulSoup(content,'html.parser'); toc=headings_toc(soup); content=str(soup)
        prev=pages[i-1] if i else None; nxt=pages[i+1] if i+1<len(pages) else None
        htmlpage=tpl.render(page_title=f'Chapter {ch.num}. {ch.title}',description=ch.subtitle,root='../',content=content,toc=toc,nav=nav_html(chapters,appmeta,ch.slug,prefix='../'),prev=prev,next=nxt)
        htmlpage=htmlpage.replace('<body>','<body data-root="../">')
        (chdir/f'{ch.slug}.html').write_text(htmlpage,encoding='utf-8')
        text=soup.get_text(' ',strip=True); heads=' '.join(h.get_text(' ',strip=True) for h in soup.find_all(['h2','h3']))
        search.append({'title':f'Chapter {ch.num}. {ch.title}','headings':heads,'text':text,'excerpt':text[:240]+'…','href':f'chapters/{ch.slug}.html'})
    # appendices
    base=len(chapters)
    for j,(title,blocks) in enumerate(appendices.items()):
        slug=slugify(title); src=(OUT/'appendices'/f'{slug}.md').read_text(encoding='utf-8'); content=render_markdown(md,src); soup=BeautifulSoup(content,'html.parser'); toc=headings_toc(soup)
        prev=pages[base+j-1] if base+j else None; nxt=pages[base+j+1] if base+j+1<len(pages) else None
        hp=tpl.render(page_title=title,description='Portable course tools and example index.',root='../',content=str(soup),toc=toc,nav=nav_html(chapters,appmeta,prefix='../'),prev=prev,next=nxt).replace('<body>','<body data-root="../">')
        (apdir/f'{slug}.html').write_text(hp,encoding='utf-8'); text=soup.get_text(' ',strip=True); search.append({'title':title,'headings':'','text':text,'excerpt':text[:240]+'…','href':f'appendices/{slug}.html'})
    # references page
    refcontent='<h1>References</h1><p>This bibliography contains only works cited in the textbook. Chapter pages also provide chapter-specific reference lists.</p>'+''.join(f'<div class="reference">{html.escape(r)}</div>' for r in all_refs)
    soup=BeautifulSoup(refcontent,'html.parser'); toc=headings_toc(soup); i=base+len(appendices)
    hp=tpl.render(page_title='References',description='Works cited in the textbook.',root='',content=str(soup),toc=toc,nav=nav,prev=pages[i-1],next=pages[i+1]).replace('<body>','<body data-root="">'); (docs/'references.html').write_text(hp,encoding='utf-8')
    search.append({'title':'References','headings':'','text':soup.get_text(' ',strip=True),'excerpt':'Complete bibliography of works cited in the textbook.','href':'references.html'})
    # about page
    report=(OUT/'STRUCTURE_AND_EDITORIAL_REPORT.md').read_text(encoding='utf-8'); content=render_markdown(md,report); soup=BeautifulSoup(content,'html.parser'); toc=headings_toc(soup)
    hp=tpl.render(page_title='About and editorial notes',description='How the revised edition is organized and sourced.',root='',content=str(soup),toc=toc,nav=nav,prev=pages[-2],next=None).replace('<body>','<body data-root="">'); (docs/'about.html').write_text(hp,encoding='utf-8')
    # homepage
    cards=[]
    for part,(pt,desc) in PARTS.items():
        cards.append(f'<section class="part-intro"><h2>Part {part}. {html.escape(pt)}</h2><p>{html.escape(desc)}</p><div class="chapter-card-grid">')
        for ch in [x for x in chapters if x.part==part]: cards.append(f'<a class="chapter-card" href="chapters/{ch.slug}.html"><small>Chapter {ch.num}</small><strong>{html.escape(ch.title)}</strong><span>{html.escape(ch.subtitle)}</span></a>')
        cards.append('</div></section>')
    home=f'''<section class="hero"><p class="eyebrow">Student textbook · 2026 edition</p><h1>Decision in the Making</h1><p>The behavioral science of choice, influence, and agreement. This edition follows the process through which decisions take shape, move between minds, and become shared action.</p><div class="download-row"><a href="downloads/Decision_Persuasion_Negotiation_Student_Ebook.epub">Download EPUB</a><a href="references.html">Browse references</a><a href="about.html">Read editorial notes</a></div></section>{''.join(cards)}'''
    hp=tpl.render(page_title='Home',description='Online textbook for Decision in the Making.',root='',content=home,toc='',nav=nav,prev=None,next=pages[0]).replace('<body>','<body data-root="">'); (docs/'index.html').write_text(hp,encoding='utf-8')
    (docs/'search-index.json').write_text(json.dumps(search,ensure_ascii=False),encoding='utf-8')

# ---------- reports/readme ----------

def write_support_files(chapters, old_refs, all_refs, removed_refs, unresolved):
    OUT.mkdir(parents=True,exist_ok=True)
    # detailed report
    lines=['# Structure and editorial report','',
    '## Why the structure changed','',
    'The original 34-chapter manuscript already had a strong modular design, but three transitions weakened the cumulative argument: the post-choice narrator appeared before the main shortcut chapters; statistical correction was separated from representativeness; and social learning, conformity, norms, authority, and influence triggers were compressed into two very broad chapters. The revised edition uses 35 shorter chapters and a clearer causal progression.','',
    '1. **Choice architecture first.** Parts I and II move from the decision loop to the rational benchmark, attention, predictive inference, valuation, expectations, heuristics, probability, bias, framing, priming, and fluency.','2. **Learning before habit.** The narrator now opens Part III, so the sequence becomes choice → explanation → repetition → habit → behavior design.','3. **The social material is decompressed.** Social learning and attribution, conformity and norms, and authority and influence triggers are now three separate chapters.','4. **Persuasion precedes story; story precedes communication.** This keeps intentional model change distinct from mutual understanding before negotiation integrates them.','5. **Negotiation remains the capstone.** It moves from joint choice to distributive preparation, bargaining tactics, integrative trade-offs, advanced agreement design, and implementation.','',
    '## Revised part structure','']
    for p,(t,d) in PARTS.items():
        lines.append(f'### Part {p}. {t}\n\n{d}\n')
        for c in [x for x in chapters if x.part==p]: lines.append(f'- Chapter {c.num}. **{c.title}** — {c.subtitle}')
    lines += ['', '## Editorial corrections','',
    '- Replaced vague or broken “another chapter” and “later chapter” references with named chapters or neutral transitions.','- Corrected “self-service” to “self-serving judgment” in the former bias subtitle.','- Removed the preface claim that the book contained four chapters.','- Replaced references to a nonexistent standalone prospect-theory chapter and a nonexistent post-negotiation nudge chapter.','- Preserved the substantive examples, activities, and editable tables while making part transitions explicit.','',
    '## Visual and table design','',
    'The revised online edition includes 21 original, accessible conceptual diagrams and retains the substantive tables from the manuscript. Figures are supplied as SVG for the web and PNG for EPUB compatibility. Tables are responsive and horizontally scroll on small screens rather than shrinking into unreadable images.','',
    '## Reference policy','',
    f'- Original master bibliography: {len(old_refs)} entries.',f'- Revised master bibliography: {len(all_refs)} unique entries cited in the text.',f'- Removed from the published bibliography because no corresponding in-text citation remained: {len(removed_refs)} entries.',f'- Unresolved author-year citations after the automated and manual audit: {len(unresolved)}.','',
    'Each chapter ends with “References cited in this chapter.” The global reference page is the deduplicated union of those chapter lists. The repository also contains a machine-readable citation audit.','',
    '## Limitations of the audit','',
    'Author-year matching is auditable but cannot determine whether every citation supports every sentence at the level of a systematic review. The audit guarantees bibliography-text correspondence, removes uncited bibliography entries, and identifies unresolved author-year strings. It does not replace expert review of every source’s full methods and boundary conditions. Claims involving contested or context-sensitive effects are written with explicit qualifications in the chapter text.']
    (OUT/'STRUCTURE_AND_EDITORIAL_REPORT.md').write_text('\n'.join(lines),encoding='utf-8')
    # citation audit
    audit={'summary':{'original_reference_entries':len(old_refs),'published_unique_cited_references':len(all_refs),'removed_uncited_entries':len(removed_refs),'unresolved_citations':len(unresolved),'chapters':len(chapters)},'chapters':[{'chapter':c.num,'title':c.title,'reference_count':len(c.references),'references':c.references} for c in chapters],'removed_uncited_references':removed_refs,'unresolved_citations':unresolved}
    (OUT/'citation-audit.json').write_text(json.dumps(audit,ensure_ascii=False,indent=2),encoding='utf-8')
    md=['# Citation audit','',f'Published references: **{len(all_refs)}**  ','',f'Removed uncited entries: **{len(removed_refs)}**  ','',f'Unresolved citations: **{len(unresolved)}**','', '## Chapter counts','', '| Chapter | References cited |','|---|---:|']
    for c in chapters: md.append(f'| {c.num}. {c.title} | {len(c.references)} |')
    md += ['','## Removed uncited entries','']+[f'- {x}' for x in removed_refs]
    if unresolved: md += ['','## Unresolved citation strings','']+[f'- {x}' for x in unresolved]
    (OUT/'CITATION_AUDIT.md').write_text('\n'.join(md),encoding='utf-8')
    # README
    readme=f'''# Decision in the Making — online textbook

This repository is ready to upload to GitHub. The student-facing website is already built in `docs/`; no local software is required to publish it.

## Publish with GitHub Pages

1. Create a repository and upload the complete contents of this folder.
2. Open **Settings → Pages**.
3. Under **Build and deployment**, either:
   - choose **Deploy from a branch**, select the default branch and `/docs`; or
   - keep **GitHub Actions** and use the included `.github/workflows/pages.yml` workflow.
4. Save. GitHub will show the public course URL after deployment.

The `docs/.nojekyll` file prevents Jekyll processing and allows the prebuilt static site to be served as-is.

## Edit and rebuild

- Student source: `chapters/*.md`, `appendices/*.md`, and `figures/*.svg`.
- Prebuilt site: `docs/`.
- The publishable site is the prebuilt `docs/` directory.
- Editable source chapters are in `chapters/`; diagrams are in `figures/`.
- `scripts/source_conversion_pipeline.py` records the one-time DOCX-to-web conversion used for this edition. It requires the original source DOCX files and is included for provenance, not for routine GitHub Pages deployment.

## Contents

- {len(chapters)} short chapters in seven parts.
- Responsive tables and 21 original diagrams with alternative text.
- Chapter-specific references plus a master bibliography containing only cited works.
- Full-text search, dark mode, mobile navigation, print styles, and previous/next navigation.
- Downloadable EPUB in `docs/downloads/`.
- Citation and structure audit reports.
'''
    (OUT/'README.md').write_text(readme,encoding='utf-8')
    (OUT/'.nojekyll').write_text('',encoding='utf-8')
    (OUT/'requirements.txt').write_text('markdown-it-py>=3.0\nJinja2>=3.1\nbeautifulsoup4>=4.12\nPyYAML>=6.0\n',encoding='utf-8')
    (OUT/'CITATION.cff').write_text('''cff-version: 1.2.0\ntitle: "Decision in the Making: The Behavioral Science of Choice, Influence, and Agreement"\nmessage: "Please cite this textbook using the metadata below."\ntype: book\nauthors:\n  - family-names: Zhang\n    given-names: Huanren Warren\nyear: 2026\nversion: "2026 GitHub edition"\n''',encoding='utf-8')
    wf=OUT/'.github/workflows'; wf.mkdir(parents=True,exist_ok=True)
    (wf/'pages.yml').write_text('''name: Deploy static textbook to GitHub Pages\non:\n  push:\n    branches: ["main", "master"]\n  workflow_dispatch:\npermissions:\n  contents: read\n  pages: write\n  id-token: write\nconcurrency:\n  group: pages\n  cancel-in-progress: true\njobs:\n  deploy:\n    environment:\n      name: github-pages\n      url: ${{ steps.deployment.outputs.page_url }}\n    runs-on: ubuntu-latest\n    steps:\n      - name: Checkout\n        uses: actions/checkout@v4\n      - name: Configure Pages\n        uses: actions/configure-pages@v5\n      - name: Upload prebuilt site\n        uses: actions/upload-pages-artifact@v3\n        with:\n          path: docs\n      - name: Deploy\n        id: deployment\n        uses: actions/deploy-pages@v4\n''',encoding='utf-8')

# ---------- unresolved audit ----------

def likely_citation_tokens(text: str):
    """Extract author-year tokens for a fast correspondence audit.

    The parser intentionally records the nearest surname in a citation. The
    resolver below accepts either a first author or a coauthor from a matching
    reference, which prevents valid expanded narrative citations from being
    mistaken for missing sources.
    """
    toks=[]
    for m in re.finditer(r'\(([^()]{0,300}\b(?:18|19|20)\d{2}[a-z]?(?:\s*/\s*(?:18|19|20)\d{2})?[^()]*)\)',text):
        chunk=m.group(1)
        for seg in chunk.split(';'):
            years=re.findall(r'\b((?:18|19|20)\d{2}[a-z]?)\b',seg)
            if not years: continue
            pre=seg[:seg.find(years[0])]
            pre=re.sub(r'^(?:e\.g\.,?|see|cf\.)\s*','',pre.strip(),flags=re.I)
            sm=re.search(r'([A-Z][A-Za-zÀ-ÖØ-öø-ÿ’\'\-]+)(?:\s+et\s+al\.|\s*(?:&|and)\s*[A-Z][A-Za-zÀ-ÖØ-öø-ÿ’\'\-]+|(?:,\s*[A-Z][A-Za-zÀ-ÖØ-öø-ÿ’\'\-]+)*)?\s*,?\s*$',pre)
            if sm:
                first=sm.group(1)
                for y in years: toks.append((first,y,seg.strip()))
    for m in re.finditer(r'\b([A-Z][A-Za-zÀ-ÖØ-öø-ÿ’\'\-]+)(?:\s+et\s+al\.|\s+(?:&|and)\s+[A-Z][A-Za-zÀ-ÖØ-öø-ÿ’\'\-]+)?(?:[’\']s)?\s*\(((?:18|19|20)\d{2}[a-z]?)\)',text):
        toks.append((m.group(1),m.group(2),m.group(0)))
    return toks


def _citation_name_variants(name: str):
    n=re.sub(r'[’\']s$','',name.strip(),flags=re.I)
    out={normalize_key(n)}
    parts=n.split()
    if parts: out.add(normalize_key(parts[-1]))
    return {x for x in out if x}


def unresolved_citations(chapters, refdb):
    # Accept first authors and coauthors. Expanded narrative citations often
    # cause a lightweight parser to land on the final coauthor; the full
    # chapter reference matcher still requires the proper complete citation.
    available=set()
    for r in refdb:
        years={r.year}
        om=re.search(r'Original work (?:published|ca\.)\s*((?:18|19|20)\d{2})',r.text,re.I)
        if om: years.add(om.group(1))
        for sn in r.surnames:
            for nv in _citation_name_variants(sn):
                for y in years: available.add((nv,y))
        # Corporate and particle names may be reduced to the last word by the parser.
        for nv in _citation_name_variants(r.first):
            for y in years: available.add((nv,y))
    un=[]
    for ch in chapters:
        text='\n'.join(b.text if isinstance(b,PBlock) else ' '.join(sum(b.rows,[])) for b in ch.blocks)
        for first,y,raw in likely_citation_tokens(text):
            if not any((nv,y) in available for nv in _citation_name_variants(first)):
                un.append({'chapter':ch.num,'first_author':first,'year':y,'citation':raw})
    seen=set(); out=[]
    for x in un:
        k=(x['chapter'],normalize_key(x['citation']))
        if k not in seen:
            seen.add(k); out.append(x)
    return out

# ---------- main ----------
def main():
    if OUT.exists(): shutil.rmtree(OUT)
    OUT.mkdir(parents=True); (OUT/'chapters').mkdir(); (OUT/'appendices').mkdir(); (OUT/'figures').mkdir(); (OUT/'scripts').mkdir()
    front,raw,appendices,old_refs=extract_book()
    chapters=restructure(front,raw)
    refdb=build_reference_db(old_refs)
    figs=generate_figures(OUT/'figures')

    # Patch front matter source into short front pages rather than repeating old TOC.
    # References per chapter are generated from in-text citations after citation insertions.
    for ch in chapters:
        body='\n'.join(b.text if isinstance(b,PBlock) else ' '.join(sum(b.rows,[])) for b in ch.blocks)
        rr=refs_for_text(body,refdb)
        ch.references=[r.text for r in rr]
        (OUT/'chapters'/f'{ch.slug}.md').write_text(blocks_to_md(ch,figs),encoding='utf-8')

    # appendices
    for title,blocks in appendices.items():
        for b in blocks:
            if isinstance(b,PBlock): b.text=patch_text(b.text)
        (OUT/'appendices'/f'{slugify(title)}.md').write_text(appendix_to_md(title,blocks),encoding='utf-8')

    all_ref_text=[]; seen=set()
    for ch in chapters:
        for r in ch.references:
            k=normalize_key(r)
            if k not in seen: seen.add(k); all_ref_text.append(r)
    all_ref_text=sorted(all_ref_text,key=lambda x:normalize_key(x.split('(')[0]))
    old_norm={normalize_key(x):x for x in old_refs}; published={normalize_key(x) for x in all_ref_text}
    removed=[old_norm[k] for k in sorted(old_norm) if k not in published]
    unresolved=unresolved_citations(chapters,refdb)

    write_support_files(chapters,old_refs,all_ref_text,removed,unresolved)
    # copy this build script as reproducible site builder, but the script expects original source; include a lighter README command note.
    shutil.copy2(__file__,OUT/'scripts/source_conversion_pipeline.py')

    # Build combined EPUB source using PNG figures and no duplicate YAML front matter.
    combined=['% Decision in the Making','% Huanren Warren Zhang','% 2026 Edition','']
    combined += ['# Preface: The Decision Point Is Usually Too Late','',
      'A visible choice is the end of a hidden process. Before anyone signs, buys, rejects, concedes, or walks away, attention has selected evidence, perception has interpreted it, expectations have shaped what seems likely, valuation has determined what matters, social cues have altered the field, and previous choices have become habits. This book begins upstream.','',
      'The revised structure follows a continuous movement from individual judgment to social influence, communication, negotiation, and decision hygiene. Each short chapter can be read independently, but the sequence is cumulative.','']
    current_part=None
    for ch in chapters:
        if ch.part!=current_part:
            current_part=ch.part; pt,pd=PARTS[current_part]; combined += [f'# Part {current_part}. {pt}','',pd,'']
        mdtxt=blocks_to_md(ch,figs,for_epub=True,include_refs=False,image_prefix='figures/')
        mdtxt=re.sub(r'^---\n.*?\n---\n','',mdtxt,flags=re.S)
        combined.append(mdtxt)
    for title,blocks in appendices.items():
        apptxt=appendix_to_md(title,blocks)
        apptxt=re.sub(r'^---\n.*?\n---\n','',apptxt,flags=re.S)
        combined.append(apptxt)
    combined += ['# References',''] + [r+'\n' for r in all_ref_text]
    (OUT/'book.md').write_text('\n'.join(combined),encoding='utf-8')

    # Build site before EPUB, then copy EPUB into downloads.
    build_site(chapters,appendices,all_ref_text,figs)
    # EPUB with embedded PNG figures.
    cmd=f"cd {OUT} && pandoc book.md -o {EPUB} --toc --toc-depth=2 --metadata title='Decision in the Making' --metadata author='Huanren Warren Zhang' --metadata lang='en'"
    rc=os.system(cmd)
    if rc!=0: raise RuntimeError('pandoc EPUB build failed')
    shutil.copy2(EPUB,OUT/'docs/downloads'/EPUB.name)

    # Rebuild homepage now that EPUB exists (same site build is fine because link target now present).
    # QA checks
    assert (OUT/'docs/index.html').exists()
    assert len(list((OUT/'docs/chapters').glob('*.html')))==len(chapters)
    assert len(list((OUT/'figures').glob('*.svg')))>=21
    # Internal href existence check.
    broken=[]
    for hp in (OUT/'docs').rglob('*.html'):
        soup=BeautifulSoup(hp.read_text(encoding='utf-8'),'html.parser')
        for a in soup.find_all('a',href=True):
            h=a['href']
            if h.startswith(('http:','https:','mailto:','#')): continue
            rel=h.split('#')[0].split('?')[0]
            if not rel: continue
            target=(hp.parent/rel).resolve()
            if not target.exists(): broken.append((str(hp.relative_to(OUT/'docs')),h))
    (OUT/'link-check.json').write_text(json.dumps({'broken':broken},indent=2),encoding='utf-8')
    if broken: print('BROKEN LINKS',broken[:20])

    # ZIP complete repo.
    if ZIP.exists(): ZIP.unlink()
    with zipfile.ZipFile(ZIP,'w',zipfile.ZIP_DEFLATED) as z:
        for f in OUT.rglob('*'):
            if f.is_file(): z.write(f,f.relative_to(OUT.parent))
    print(json.dumps({'out':str(OUT),'zip':str(ZIP),'epub':str(EPUB),'chapters':len(chapters),'references':len(all_ref_text),'removed':len(removed),'unresolved':len(unresolved),'broken_links':len(broken)},indent=2))

if __name__=='__main__': main()
